
import json
import re
import sqlite3
import uuid
from contextlib import contextmanager
from datetime import date, datetime
from pathlib import Path
from typing import Any, Dict, Iterable, Optional, List
from io import BytesIO
from html import escape

import pandas as pd
import streamlit as st
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Cm
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_LEFT
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import cm, mm
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak


APP_TITLE = "SIR · Módulo 8 · Casos"
DB_PATH = Path(__file__).with_name("sir_consultas_quejas.db")
UPLOAD_DIR = Path(__file__).with_name("archivos_consultas_quejas")
UPLOAD_DIR.mkdir(exist_ok=True)

# Los datos de demostración quedan desactivados por defecto.
ENABLE_DEMO_DATA = False


# =========================================================
# CONFIGURACIÓN Y UTILIDADES
# =========================================================

st.set_page_config(
    page_title=APP_TITLE,
    page_icon="📋",
    layout="wide",
    initial_sidebar_state="expanded",
)


@contextmanager
def db_connection():
    conn = sqlite3.connect(DB_PATH)
    conn.row_factory = sqlite3.Row
    conn.execute("PRAGMA foreign_keys = ON")
    try:
        yield conn
        conn.commit()
    except Exception:
        conn.rollback()
        raise
    finally:
        conn.close()


def now_iso() -> str:
    return datetime.now().replace(microsecond=0).isoformat(sep=" ")


def normalize_text(value: Optional[str]) -> str:
    return (value or "").strip()


def json_safe(value: Any) -> Any:
    if isinstance(value, (date, datetime)):
        return value.isoformat()
    return value


def generate_id(prefix: str) -> str:
    return f"{prefix}-{uuid.uuid4().hex[:12].upper()}"


def generate_case_code(conn: sqlite3.Connection, year: int) -> str:
    row = conn.execute(
        """
        SELECT COUNT(*) AS total
        FROM casos
        WHERE substr(fecha_recepcion, 1, 4) = ?
        """,
        (str(year),),
    ).fetchone()
    consecutive = int(row["total"]) + 1
    return f"CQ-{year}-{consecutive:05d}"


def rows_to_df(rows: Iterable[sqlite3.Row]) -> pd.DataFrame:
    return pd.DataFrame([dict(r) for r in rows])


def save_uploaded_files(case_id: str, files, document_type: str, user_name: str) -> None:
    if not files:
        return

    case_dir = UPLOAD_DIR / case_id
    case_dir.mkdir(exist_ok=True)

    with db_connection() as conn:
        for uploaded in files:
            safe_name = Path(uploaded.name).name
            target = case_dir / f"{uuid.uuid4().hex[:8]}_{safe_name}"
            target.write_bytes(uploaded.getbuffer())

            conn.execute(
                """
                INSERT INTO documentos (
                    id_documento, id_caso, tipo_documento, nombre_archivo,
                    ruta_archivo, fecha_carga, cargado_por
                ) VALUES (?, ?, ?, ?, ?, ?, ?)
                """,
                (
                    generate_id("DOC"),
                    case_id,
                    document_type,
                    safe_name,
                    str(target),
                    now_iso(),
                    user_name,
                ),
            )


def audit_change(
    conn: sqlite3.Connection,
    case_id: str,
    action: str,
    user_name: str,
    previous_data: Optional[Dict[str, Any]] = None,
    new_data: Optional[Dict[str, Any]] = None,
    details: str = "",
) -> None:
    conn.execute(
        """
        INSERT INTO auditoria (
            id_auditoria, id_caso, accion, usuario, fecha_hora,
            datos_anteriores, datos_nuevos, detalle
        ) VALUES (?, ?, ?, ?, ?, ?, ?, ?)
        """,
        (
            generate_id("AUD"),
            case_id,
            action,
            user_name,
            now_iso(),
            json.dumps(previous_data or {}, ensure_ascii=False, default=json_safe),
            json.dumps(new_data or {}, ensure_ascii=False, default=json_safe),
            details,
        ),
    )


def register_state_change(
    conn: sqlite3.Connection,
    case_id: str,
    previous_state: Optional[str],
    new_state: str,
    user_name: str,
    reason: str,
    source_followup_id: Optional[str] = None,
) -> None:
    if previous_state == new_state:
        return

    conn.execute(
        """
        INSERT INTO historial_estados (
            id_historial, id_caso, estado_anterior, estado_nuevo,
            fecha_cambio, usuario_cambio, motivo, id_seguimiento_origen
        ) VALUES (?, ?, ?, ?, ?, ?, ?, ?)
        """,
        (
            generate_id("EST"),
            case_id,
            previous_state,
            new_state,
            now_iso(),
            user_name,
            reason,
            source_followup_id,
        ),
    )


def get_case(conn: sqlite3.Connection, case_id: str) -> Optional[sqlite3.Row]:
    return conn.execute("SELECT * FROM casos WHERE id_caso = ?", (case_id,)).fetchone()


def update_case_fields(
    conn: sqlite3.Connection,
    case_id: str,
    changes: Dict[str, Any],
    user_name: str,
    action: str,
    detail: str = "",
) -> None:
    current = get_case(conn, case_id)
    if not current:
        raise ValueError("El caso no existe.")

    previous = {key: current[key] for key in changes.keys()}
    cleaned = {key: json_safe(value) for key, value in changes.items()}

    assignments = ", ".join(f"{field} = ?" for field in cleaned.keys())
    values = list(cleaned.values()) + [now_iso(), user_name, case_id]

    conn.execute(
        f"""
        UPDATE casos
        SET {assignments},
            fecha_ultima_actualizacion = ?,
            actualizado_por = ?
        WHERE id_caso = ?
        """,
        values,
    )

    audit_change(
        conn,
        case_id=case_id,
        action=action,
        user_name=user_name,
        previous_data=previous,
        new_data=cleaned,
        details=detail,
    )


def required_text(value: str, label: str) -> None:
    if not normalize_text(value):
        raise ValueError(f"El campo '{label}' es obligatorio.")


def validate_case_dates(
    fecha_recepcion: date,
    fecha_registro: date,
    fecha_asignacion: Optional[date] = None,
    fecha_cierre: Optional[date] = None,
) -> None:
    if fecha_registro < fecha_recepcion:
        raise ValueError("La fecha de registro no puede ser anterior a la fecha de recepción.")
    if fecha_asignacion and fecha_asignacion < fecha_registro:
        raise ValueError("La fecha de asignación no puede ser anterior a la fecha de registro.")
    # Se permite abrir y cerrar el mismo día.
    if fecha_cierre and fecha_cierre < fecha_registro:
        raise ValueError("La fecha de cierre no puede ser anterior a la fecha de registro.")


def validate_required_case_location(payload: Dict[str, Any]) -> None:
    """La ubicación es obligatoria cuando el caso está clasificado como Queja."""
    if payload.get("clasificacion") != "Queja":
        return
    required = {
        "provincia_hecho": "Provincia del caso",
        "distrito_hecho": "Distrito del caso",
        "corregimiento_hecho": "Corregimiento del caso",
        "lugar_poblado_hecho": "Comunidad del caso",
        "direccion_hecho": "Dirección o ubicación exacta del caso",
    }
    missing = [label for field, label in required.items() if not normalize_text(str(payload.get(field) or ""))]
    if missing:
        raise ValueError(
            "Para un caso clasificado como Queja debe completar la ubicación del caso: "
            + ", ".join(missing)
            + "."
        )



# =========================================================
# IMPORTACIÓN SURVEY123
# =========================================================

def _xtext(value):
    if pd.isna(value):
        return ""
    return str(value).strip()

def _xdate(value):
    if pd.isna(value):
        return None
    parsed = pd.to_datetime(value, errors="coerce")
    return None if pd.isna(parsed) else parsed.date().isoformat()

def _xtime(value):
    if pd.isna(value):
        return None
    parsed = pd.to_datetime(value, errors="coerce")
    return None if pd.isna(parsed) else parsed.strftime("%H:%M")

def _xbool(value):
    return int(_xtext(value).lower() in {"si", "sí", "yes", "true", "1", "x"})


def _xcode(value: Any, width: int) -> str:
    text = _xtext(value)
    if not text:
        return ""
    if text.endswith(".0") and text[:-2].isdigit():
        text = text[:-2]
    if text == "63":
        return text
    return text.zfill(width) if text.isdigit() and len(text) <= width else text


def _survey_status(value):
    text = _xtext(value).lower()
    if text in {"cerrada", "cerrado"}:
        return "Cerrada"
    if text in {"abierta", "abierto"}:
        return "En atención"
    return "Registrada"


SURVEY123_EXPECTED_FIELDS = {
    "Quejas_y_Consultas_Rio_Indi_0": [
        "ObjectID", "GlobalID", "Fecha de registro", "Medio por el que se recibe",
        "Otro medio de recepción", "Nombre del contacto", "Sexo del contacto",
        "Cédula del contacto", "Provincia del contacto (cod.)",
        "Distrito del contacto (cod.)", "Corregimiento del contacto (cod.)",
        "Comunidad del contacto (cod.)", "Provincia del contacto",
        "Distrito del contacto", "Corregimiento del contacto", "Comunidad del contacto",
        "Dirección del contacto", "Recepción de Consulta o Queja por:",
        "Contacto de la persona", "Teléfono o celular", "Teléfono del contacto",
        "Celular del contacto", "Correo electrónico del contacto", "Clasificación",
        "N° del formulario", "Tema de la consulta o queja", "Tipo de queja", "¿Cuál?",
        "Responsable del origen de la queja", "Nombre del responsable del origen de la queja",
        "Descripción consulta o queja", "¿Ha sido presentada anteriormente?",
        "Respuesta inmediata a la queja", "Propuesta mejor solución queja",
        "Equipo Supervisor", "Supervisor de Equipo de Gestión Socioambiental (PH)",
        "Fecha de asignación", "Cantidad de responsables de la asignación",
        "Responsable 1", "Equipo Responsable 1", "Responsable 2", "Equipo Responsable 2",
        "Responsable 3", "Equipo Responsable 3", "seccionSupervisor",
        "seccionResponsable1", "seccionResponsable2", "seccionResponsable3", "Estatus",
        "Fecha de cierre", "Respuesta brindada", "Aceptación contacto",
        "Comentarios de aceptación contacto", "Comentarios de rechazo contacto",
        "CreationDate", "Creator", "EditDate", "Editor", "Satisfacción contacto",
        "Se brindó una respuesta", "x", "y",
    ],
    "ubicacionQueja_1": [
        "ObjectID", "GlobalID", "Provincia de la queja (cod.)",
        "Distrito de la queja (cod.)", "Corregimiento de la queja (cod.)",
        "Comunidad de la queja (cod.)", "Provincia de la queja", "Distrito de la queja",
        "Corregimiento de la queja", "Comunidad de la queja", "Dirección de la queja",
        "Referencia de la afectación", "Otra referencia de la afectación",
        "Llegada al lugar de afectación", "Coordenadas geográficas del sitio",
        "Coordenadas Norte UTM", "Coordenadas Este UTM", "numFormularioUQ",
        "ParentGlobalID", "CreationDate", "Creator", "EditDate", "Editor", "x", "y",
    ],
    "seguimiento_2": [
        "ObjectID", "GlobalID", "Fecha", "Descripción y seguimiento de la",
        "clasificacionS", "numFormularioS", "ParentGlobalID", "CreationDate",
        "Creator", "EditDate", "Editor",
    ],
}


def _jsonable(value: Any) -> Any:
    if value is None:
        return None
    try:
        if pd.isna(value):
            return None
    except (TypeError, ValueError):
        pass
    if isinstance(value, (date, datetime, pd.Timestamp)):
        return value.isoformat()
    if hasattr(value, "item"):
        try:
            return value.item()
        except Exception:
            pass
    return value


def _row_json(row: pd.Series) -> str:
    return json.dumps(
        {str(key): _jsonable(value) for key, value in row.to_dict().items()},
        ensure_ascii=False,
        default=json_safe,
    )


def _source_filename(uploaded_file: Any, default: str = "survey123.xlsx") -> str:
    return Path(getattr(uploaded_file, "name", default)).name


def _next_control_code(conn: sqlite3.Connection, table: str, column: str, prefix: str) -> str:
    rows = conn.execute(
        f'SELECT "{column}" AS code FROM "{table}" '
        f'WHERE "{column}" LIKE ? ORDER BY rowid',
        (f"{prefix}%",),
    ).fetchall()
    maximum = 0
    for row in rows:
        code = str(row["code"] or "")
        match = re.search(r"(\d+)$", code)
        if match:
            maximum = max(maximum, int(match.group(1)))
    return f"{prefix}{maximum + 1:06d}"


def _backfill_control_codes(
    conn: sqlite3.Connection,
    table: str,
    pk: str,
    column: str,
    prefix: str,
) -> None:
    rows = conn.execute(
        f'SELECT "{pk}" AS pk FROM "{table}" '
        f'WHERE "{column}" IS NULL OR TRIM("{column}") = "" ORDER BY rowid'
    ).fetchall()
    for row in rows:
        code = _next_control_code(conn, table, column, prefix)
        conn.execute(
            f'UPDATE "{table}" SET "{column}" = ? WHERE "{pk}" = ?',
            (code, row["pk"]),
        )


def _assign_control_code(
    conn: sqlite3.Connection,
    table: str,
    pk: str,
    pk_value: str,
    column: str,
    prefix: str,
) -> str:
    row = conn.execute(
        f'SELECT "{column}" AS code FROM "{table}" WHERE "{pk}" = ?',
        (pk_value,),
    ).fetchone()
    if row and normalize_text(row["code"]):
        return row["code"]
    code = _next_control_code(conn, table, column, prefix)
    conn.execute(
        f'UPDATE "{table}" SET "{column}" = ? WHERE "{pk}" = ?',
        (code, pk_value),
    )
    return code


def _upsert_survey_reference(
    conn: sqlite3.Connection,
    tipo_registro: str,
    id_caso: str,
    id_registro_relacionado: Optional[str],
    survey_objectid: str,
    survey_globalid: str,
    survey_parentglobalid: str,
    numero_formulario: str,
    archivo_origen: str,
) -> None:
    if not any([survey_objectid, survey_globalid, survey_parentglobalid, numero_formulario]):
        return
    existing = None
    if survey_globalid:
        existing = conn.execute(
            """
            SELECT id_referencia FROM referencias_survey123
            WHERE tipo_registro = ? AND survey_globalid = ?
            """,
            (tipo_registro, survey_globalid),
        ).fetchone()
    values = {
        "tipo_registro": tipo_registro,
        "id_caso": id_caso,
        "id_registro_relacionado": id_registro_relacionado,
        "survey_objectid": survey_objectid,
        "survey_globalid": survey_globalid,
        "survey_parentglobalid": survey_parentglobalid,
        "numero_formulario": numero_formulario,
        "archivo_origen": archivo_origen,
        "fecha_importacion": now_iso(),
    }
    if existing:
        _update(conn, "referencias_survey123", "id_referencia", existing["id_referencia"], values)
    else:
        columns = ["id_referencia"] + list(values.keys())
        conn.execute(
            f"INSERT INTO referencias_survey123 ({', '.join(columns)}) "
            f"VALUES ({', '.join(['?'] * len(columns))})",
            [generate_id("REFS123")] + list(values.values()),
        )


def _save_pending_survey_record(
    conn: sqlite3.Connection,
    tipo_registro: str,
    row: pd.Series,
    objectid: str,
    globalid: str,
    parentglobalid: str,
    form_number: str,
    source_file: str,
    reason: str,
) -> None:
    existing = None
    if globalid:
        existing = conn.execute(
            """
            SELECT id_pendiente FROM registros_survey_pendientes
            WHERE tipo_registro = ? AND survey_globalid = ?
            ORDER BY fecha_importacion DESC LIMIT 1
            """,
            (tipo_registro, globalid),
        ).fetchone()
    values = {
        "tipo_registro": tipo_registro,
        "survey_objectid": objectid,
        "survey_globalid": globalid,
        "survey_parentglobalid": parentglobalid,
        "numero_formulario": form_number,
        "datos_originales": _row_json(row),
        "archivo_origen": source_file,
        "motivo": reason,
        "estado": "Pendiente",
        "fecha_importacion": now_iso(),
        "id_caso_resuelto": None,
        "fecha_resolucion": None,
    }
    if existing:
        _update(
            conn, "registros_survey_pendientes", "id_pendiente",
            existing["id_pendiente"], values,
        )
    else:
        columns = ["id_pendiente"] + list(values.keys())
        conn.execute(
            f"INSERT INTO registros_survey_pendientes ({', '.join(columns)}) "
            f"VALUES ({', '.join(['?'] * len(columns))})",
            [generate_id("PENDS123")] + list(values.values()),
        )


def survey_structure_report(uploaded_file: Any) -> pd.DataFrame:
    """Compara las hojas y encabezados recibidos con la estructura esperada."""
    uploaded_file.seek(0)
    xls = pd.ExcelFile(uploaded_file)
    records = []
    for sheet_name, expected in SURVEY123_EXPECTED_FIELDS.items():
        if sheet_name not in xls.sheet_names:
            records.append({
                "hoja": sheet_name,
                "estado": "Hoja faltante",
                "campos_esperados": len(expected),
                "campos_detectados": 0,
                "faltantes": ", ".join(expected),
                "adicionales": "",
            })
            continue
        detected = pd.read_excel(xls, sheet_name=sheet_name, nrows=0).columns.astype(str).tolist()
        missing = [field for field in expected if field not in detected]
        extra = [field for field in detected if field not in expected]
        records.append({
            "hoja": sheet_name,
            "estado": "Completa" if not missing else "Revisar",
            "campos_esperados": len(expected),
            "campos_detectados": len(detected),
            "faltantes": ", ".join(missing),
            "adicionales": ", ".join(extra),
        })
    uploaded_file.seek(0)
    return pd.DataFrame(records)


def ensure_import_schema(conn):
    def add_column(table, name, definition):
        cols = {r["name"] for r in conn.execute(f"PRAGMA table_info({table})").fetchall()}
        if name not in cols:
            conn.execute(f'ALTER TABLE "{table}" ADD COLUMN "{name}" {definition}')

    # Campos internos, campos exactos de Survey123 y metadatos del caso.
    case_columns = {
        "id_control_m8": "TEXT",
        "provincia_contacto_codigo": "TEXT",
        "distrito_contacto_codigo": "TEXT",
        "corregimiento_contacto_codigo": "TEXT",
        "lugar_poblado_contacto_codigo": "TEXT",
        "provincia_hecho_codigo": "TEXT",
        "distrito_hecho_codigo": "TEXT",
        "corregimiento_hecho_codigo": "TEXT",
        "lugar_poblado_hecho_codigo": "TEXT",
        "contacto_persona": "TEXT",
        "telefono_o_celular": "TEXT",
        "tipo_caso": "TEXT",
        "tipo_caso_otro": "TEXT",
        "tipo_queja_otro": "TEXT",  # compatibilidad histórica
        "responsable_origen_catalogo": "TEXT",
        "nombre_responsable_origen": "TEXT",
        "equipo_supervisor": "TEXT",
        "cantidad_responsables_asignacion": "INTEGER",
        "equipo_responsable_1": "TEXT",
        "equipo_responsable_2": "TEXT",
        "equipo_responsable_3": "TEXT",
        "seccion_supervisor": "TEXT",
        "seccion_responsable_1": "TEXT",
        "seccion_responsable_2": "TEXT",
        "seccion_responsable_3": "TEXT",
        "comentarios_aceptacion_contacto": "TEXT",
        "comentarios_rechazo_contacto": "TEXT",
        "se_brindo_respuesta": "TEXT",
        "hora_cierre": "TEXT",
        "hora_comunicacion_cierre": "TEXT",
        "autoriza_divulgacion_datos": "TEXT",
        "referencia_afectacion": "TEXT",
        "otra_referencia_afectacion": "TEXT",
        "llegada_lugar_afectacion": "TEXT",
        "tiene_coordenadas_hecho": "TEXT",
        "coordenada_norte_utm": "REAL",
        "coordenada_este_utm": "REAL",
        "fuente_registro": "TEXT",
        "archivo_origen": "TEXT",
        "survey_globalid": "TEXT",
        "survey_objectid": "TEXT",
        "survey_creationdate": "TEXT",
        "survey_editdate": "TEXT",
        "survey_creator": "TEXT",
        "survey_editor": "TEXT",
        "survey_datos_originales": "TEXT",
    }
    for name, definition in case_columns.items():
        add_column("casos", name, definition)

    followup_columns = {
        "id_control_seguimiento": "TEXT",
        "clasificacion_survey": "TEXT",
        "numero_formulario_survey": "TEXT",
        "archivo_origen": "TEXT",
        "survey_globalid": "TEXT",
        "survey_parentglobalid": "TEXT",
        "survey_objectid": "TEXT",
        "survey_creationdate": "TEXT",
        "survey_editdate": "TEXT",
        "survey_creator": "TEXT",
        "survey_editor": "TEXT",
        "survey_datos_originales": "TEXT",
        "informo_solicitante": "INTEGER NOT NULL DEFAULT 0",
        "tipo_informacion_solicitante": "TEXT",
        "medio_informacion_solicitante": "TEXT",
        "resultado_contacto_solicitante": "TEXT",
        "contenido_comunicado": "TEXT",
    }
    for name, definition in followup_columns.items():
        add_column("seguimientos", name, definition)

    conn.executescript("""
        CREATE TABLE IF NOT EXISTS ubicaciones_survey123 (
            id_ubicacion TEXT PRIMARY KEY,
            id_caso TEXT NOT NULL,
            id_control_ubicacion TEXT,
            survey_globalid TEXT,
            survey_parentglobalid TEXT,
            survey_objectid TEXT,
            numero_formulario TEXT,
            provincia_codigo TEXT,
            distrito_codigo TEXT,
            corregimiento_codigo TEXT,
            comunidad_codigo TEXT,
            provincia TEXT,
            distrito TEXT,
            corregimiento TEXT,
            comunidad TEXT,
            direccion TEXT,
            referencia_afectacion TEXT,
            otra_referencia TEXT,
            llegada_lugar TEXT,
            coordenadas_geograficas TEXT,
            coordenada_norte_utm REAL,
            coordenada_este_utm REAL,
            latitud REAL,
            longitud REAL,
            survey_creationdate TEXT,
            survey_editdate TEXT,
            survey_creator TEXT,
            survey_editor TEXT,
            survey_datos_originales TEXT,
            archivo_origen TEXT,
            fecha_importacion TEXT NOT NULL,
            FOREIGN KEY(id_caso) REFERENCES casos(id_caso)
        );
        CREATE TABLE IF NOT EXISTS importaciones_survey123 (
            id_importacion TEXT PRIMARY KEY, nombre_archivo TEXT NOT NULL,
            fecha_importacion TEXT NOT NULL, usuario TEXT NOT NULL,
            casos_nuevos INTEGER DEFAULT 0, casos_actualizados INTEGER DEFAULT 0,
            casos_sin_cambios INTEGER DEFAULT 0, ubicaciones_nuevas INTEGER DEFAULT 0,
            ubicaciones_actualizadas INTEGER DEFAULT 0, ubicaciones_sin_cambios INTEGER DEFAULT 0,
            seguimientos_nuevos INTEGER DEFAULT 0, seguimientos_actualizados INTEGER DEFAULT 0,
            seguimientos_sin_cambios INTEGER DEFAULT 0, errores INTEGER DEFAULT 0,
            detalle_errores TEXT
        );
        CREATE TABLE IF NOT EXISTS referencias_survey123 (
            id_referencia TEXT PRIMARY KEY,
            tipo_registro TEXT NOT NULL,
            id_caso TEXT NOT NULL,
            id_registro_relacionado TEXT,
            survey_objectid TEXT,
            survey_globalid TEXT,
            survey_parentglobalid TEXT,
            numero_formulario TEXT,
            archivo_origen TEXT,
            fecha_importacion TEXT NOT NULL,
            FOREIGN KEY(id_caso) REFERENCES casos(id_caso)
        );
        CREATE TABLE IF NOT EXISTS registros_survey_pendientes (
            id_pendiente TEXT PRIMARY KEY,
            tipo_registro TEXT NOT NULL,
            survey_objectid TEXT,
            survey_globalid TEXT,
            survey_parentglobalid TEXT,
            numero_formulario TEXT,
            datos_originales TEXT NOT NULL,
            archivo_origen TEXT,
            motivo TEXT,
            estado TEXT NOT NULL DEFAULT 'Pendiente',
            fecha_importacion TEXT NOT NULL,
            id_caso_resuelto TEXT,
            fecha_resolucion TEXT
        );
    """)

    add_column("importaciones_survey123", "pendientes_vinculacion", "INTEGER DEFAULT 0")

    location_columns = {
        "id_control_ubicacion": "TEXT",
        "provincia_codigo": "TEXT",
        "distrito_codigo": "TEXT",
        "corregimiento_codigo": "TEXT",
        "comunidad_codigo": "TEXT",
        "survey_datos_originales": "TEXT",
        "archivo_origen": "TEXT",
    }
    for name, definition in location_columns.items():
        add_column("ubicaciones_survey123", name, definition)

    # Los GlobalID se conservan para intercambio futuro, pero no funcionan como ID interno.
    conn.executescript("""
        CREATE INDEX IF NOT EXISTS idx_casos_survey_gid ON casos(survey_globalid);
        CREATE INDEX IF NOT EXISTS idx_seg_survey_gid ON seguimientos(survey_globalid);
        CREATE INDEX IF NOT EXISTS idx_ubi_survey_gid ON ubicaciones_survey123(survey_globalid);
        CREATE INDEX IF NOT EXISTS idx_ref_survey_gid ON referencias_survey123(tipo_registro, survey_globalid);
        CREATE INDEX IF NOT EXISTS idx_pend_survey_gid ON registros_survey_pendientes(tipo_registro, survey_globalid);
        CREATE UNIQUE INDEX IF NOT EXISTS idx_casos_control_m8 ON casos(id_control_m8)
            WHERE id_control_m8 IS NOT NULL AND id_control_m8 <> '';
        CREATE UNIQUE INDEX IF NOT EXISTS idx_seg_control_m8 ON seguimientos(id_control_seguimiento)
            WHERE id_control_seguimiento IS NOT NULL AND id_control_seguimiento <> '';
        CREATE UNIQUE INDEX IF NOT EXISTS idx_ubi_control_m8 ON ubicaciones_survey123(id_control_ubicacion)
            WHERE id_control_ubicacion IS NOT NULL AND id_control_ubicacion <> '';
    """)

    _backfill_control_codes(conn, "casos", "id_caso", "id_control_m8", "M8-C-")
    _backfill_control_codes(
        conn, "seguimientos", "id_seguimiento", "id_control_seguimiento", "M8-S-"
    )
    _backfill_control_codes(
        conn, "ubicaciones_survey123", "id_ubicacion", "id_control_ubicacion", "M8-U-"
    )

def _changed(existing, values):
    return any(("" if existing[k] is None else str(existing[k])) != ("" if v is None else str(v)) for k,v in values.items())

def _update(conn, table, pk, pk_value, values):
    assignments = ", ".join(f'"{k}"=?' for k in values)
    conn.execute(f'UPDATE "{table}" SET {assignments} WHERE "{pk}"=?', list(values.values())+[pk_value])

def _parent_case(conn, parent_gid, form_number):
    if parent_gid:
        row = conn.execute(
            "SELECT * FROM casos WHERE survey_globalid = ?", (parent_gid,)
        ).fetchone()
        if row:
            return row
        ref = conn.execute(
            """
            SELECT c.*
            FROM referencias_survey123 r
            JOIN casos c ON c.id_caso = r.id_caso
            WHERE r.survey_globalid = ?
            ORDER BY r.fecha_importacion DESC LIMIT 1
            """,
            (parent_gid,),
        ).fetchone()
        if ref:
            return ref
    if form_number:
        return conn.execute(
            "SELECT * FROM casos WHERE codigo_caso = ?", (form_number,)
        ).fetchone()
    return None


def _existing_source_record(
    conn: sqlite3.Connection,
    table: str,
    globalid: str,
    objectid: str,
    form_column: str,
    form_number: str,
):
    if globalid:
        row = conn.execute(
            f'SELECT * FROM "{table}" WHERE survey_globalid = ?', (globalid,)
        ).fetchone()
        if row:
            return row
    if objectid and form_number:
        row = conn.execute(
            f'SELECT * FROM "{table}" WHERE survey_objectid = ? AND "{form_column}" = ?',
            (objectid, form_number),
        ).fetchone()
        if row:
            return row
    return None


def importar_survey123(uploaded_file, user_name):
    required = set(SURVEY123_EXPECTED_FIELDS)
    uploaded_file.seek(0)
    xls = pd.ExcelFile(uploaded_file)
    missing = required - set(xls.sheet_names)
    if missing:
        raise ValueError("Faltan hojas: " + ", ".join(sorted(missing)))

    casos_df = pd.read_excel(xls, sheet_name="Quejas_y_Consultas_Rio_Indi_0", dtype=object)
    ubi_df = pd.read_excel(xls, sheet_name="ubicacionQueja_1", dtype=object)
    seg_df = pd.read_excel(xls, sheet_name="seguimiento_2", dtype=object)
    source_file = _source_filename(uploaded_file)
    counters = {key: 0 for key in [
        "casos_nuevos", "casos_actualizados", "casos_sin_cambios",
        "ubicaciones_nuevas", "ubicaciones_actualizadas", "ubicaciones_sin_cambios",
        "seguimientos_nuevos", "seguimientos_actualizados", "seguimientos_sin_cambios",
        "pendientes_vinculacion", "errores",
    ]}
    errores = []

    with db_connection() as conn:
        ensure_import_schema(conn)

        for index, row in casos_df.iterrows():
            try:
                gid = _xtext(row.get("GlobalID"))
                objectid = _xtext(row.get("ObjectID"))
                form_number = _xtext(row.get("N° del formulario"))
                fecha_registro = _xdate(row.get("Fecha de registro")) or date.today().isoformat()
                hora_registro = _xtime(row.get("Fecha de registro")) or "08:00"
                clasificacion = (
                    "Queja" if "queja" in _xtext(row.get("Clasificación")).lower()
                    else "Consulta"
                )
                nombre = _xtext(row.get("Nombre del contacto"))
                cedula = _xtext(row.get("Cédula del contacto"))
                responsable_catalogo = _xtext(row.get("Responsable del origen de la queja"))
                nombre_responsable = _xtext(row.get("Nombre del responsable del origen de la queja"))
                existing = None
                matched_by_globalid = False
                if gid:
                    existing = conn.execute(
                        "SELECT * FROM casos WHERE survey_globalid = ?", (gid,)
                    ).fetchone()
                    matched_by_globalid = existing is not None
                if not existing and form_number:
                    existing = conn.execute(
                        "SELECT * FROM casos WHERE codigo_caso = ?", (form_number,)
                    ).fetchone()
                additional_survey_reference = bool(
                    existing
                    and gid
                    and _xtext(existing["survey_globalid"])
                    and gid != _xtext(existing["survey_globalid"])
                    and not matched_by_globalid
                )

                values = {
                    "codigo_caso": form_number or (
                        existing["codigo_caso"] if existing else generate_case_code(conn, int(fecha_registro[:4]))
                    ),
                    "fecha_recepcion": fecha_registro,
                    "hora_recepcion": hora_registro,
                    "fecha_registro": fecha_registro,
                    "hora_registro": hora_registro,
                    "registrado_por": _xtext(row.get("Creator")) or user_name,
                    "clasificacion": clasificacion,
                    "estado_actual": _survey_status(row.get("Estatus")),
                    "situacion_atencion": "",
                    "dentro_alcance": 1,
                    "motivo_fuera_alcance": "",
                    "tipo_identificacion_solicitante": "Identificado",
                    "nombre_solicitante": nombre,
                    "sexo": _xtext(row.get("Sexo del contacto")),
                    "cedula": cedula,
                    "telefono": _xtext(row.get("Teléfono del contacto")),
                    "celular": _xtext(row.get("Celular del contacto")),
                    "correo": _xtext(row.get("Correo electrónico del contacto")),
                    "provincia_contacto_codigo": _xcode(row.get("Provincia del contacto (cod.)"), 2),
                    "distrito_contacto_codigo": _xcode(row.get("Distrito del contacto (cod.)"), 4),
                    "corregimiento_contacto_codigo": _xcode(row.get("Corregimiento del contacto (cod.)"), 6),
                    "lugar_poblado_contacto_codigo": _xcode(row.get("Comunidad del contacto (cod.)"), 9),
                    "provincia_contacto": _xtext(row.get("Provincia del contacto")),
                    "distrito_contacto": _xtext(row.get("Distrito del contacto")),
                    "corregimiento_contacto": _xtext(row.get("Corregimiento del contacto")),
                    "lugar_poblado_contacto": _xtext(row.get("Comunidad del contacto")),
                    "direccion_contacto": _xtext(row.get("Dirección del contacto")),
                    "medio_recepcion": _xtext(row.get("Medio por el que se recibe")) or "Otro",
                    "otro_medio_recepcion": _xtext(row.get("Otro medio de recepción")),
                    "recibido_por": _xtext(row.get("Recepción de Consulta o Queja por:")) or user_name,
                    "contacto_persona": _xtext(row.get("Contacto de la persona")),
                    "telefono_o_celular": _xtext(row.get("Teléfono o celular")),
                    "tema": _xtext(row.get("Tema de la consulta o queja")),
                    "tipo_caso": _xtext(row.get("Tipo de queja")),
                    "tipo_caso_otro": _xtext(row.get("¿Cuál?")),
                    "tipo_queja": _xtext(row.get("Tipo de queja")),  # compatibilidad
                    "tipo_queja_otro": _xtext(row.get("¿Cuál?")),  # compatibilidad
                    "responsable_origen_catalogo": responsable_catalogo,
                    "nombre_responsable_origen": nombre_responsable,
                    "responsable_origen": nombre_responsable or responsable_catalogo,
                    "descripcion": _xtext(row.get("Descripción consulta o queja")) or "Sin descripción",
                    "presentado_anteriormente": _xbool(row.get("¿Ha sido presentada anteriormente?")),
                    "referencia_caso_anterior": "",
                    "respuesta_inmediata": _xtext(row.get("Respuesta inmediata a la queja")),
                    "propuesta_solucion": _xtext(row.get("Propuesta mejor solución queja")),
                    "equipo_supervisor": _xtext(row.get("Equipo Supervisor")),
                    "supervisor": _xtext(row.get("Supervisor de Equipo de Gestión Socioambiental (PH)")),
                    "fecha_asignacion": _xdate(row.get("Fecha de asignación")),
                    "cantidad_responsables_asignacion": int(row.get("Cantidad de responsables de la asignación"))
                        if not pd.isna(row.get("Cantidad de responsables de la asignación")) else None,
                    "responsable_principal": _xtext(row.get("Responsable 1")),
                    "equipo_responsable_1": _xtext(row.get("Equipo Responsable 1")),
                    "responsable_apoyo_1": _xtext(row.get("Responsable 2")),
                    "equipo_responsable_2": _xtext(row.get("Equipo Responsable 2")),
                    "responsable_apoyo_2": _xtext(row.get("Responsable 3")),
                    "equipo_responsable_3": _xtext(row.get("Equipo Responsable 3")),
                    "seccion_supervisor": _xtext(row.get("seccionSupervisor")),
                    "seccion_responsable_1": _xtext(row.get("seccionResponsable1")),
                    "seccion_responsable_2": _xtext(row.get("seccionResponsable2")),
                    "seccion_responsable_3": _xtext(row.get("seccionResponsable3")),
                    "asignado_por": _xtext(row.get("Editor")) or user_name,
                    "motivo_asignacion": "",
                    "respuesta_final": _xtext(row.get("Respuesta brindada")),
                    "acepta_respuesta": _xtext(row.get("Aceptación contacto")),
                    "comentarios_aceptacion_contacto": _xtext(row.get("Comentarios de aceptación contacto")),
                    "comentarios_rechazo_contacto": _xtext(row.get("Comentarios de rechazo contacto")),
                    "comentarios_finales": _xtext(row.get("Comentarios de aceptación contacto")),
                    "comentario_insatisfaccion": _xtext(row.get("Comentarios de rechazo contacto")),
                    "nivel_satisfaccion": _xtext(row.get("Satisfacción contacto")),
                    "se_brindo_respuesta": _xtext(row.get("Se brindó una respuesta")),
                    "fecha_cierre": _xdate(row.get("Fecha de cierre")),
                    "cerrado_por": _xtext(row.get("Editor")),
                    "motivo_cierre": "",
                    "longitud": None if pd.isna(row.get("x")) else row.get("x"),
                    "latitud": None if pd.isna(row.get("y")) else row.get("y"),
                    "fuente_registro": "Survey123",
                    "archivo_origen": source_file,
                    "survey_objectid": objectid,
                    "survey_creationdate": _xtext(row.get("CreationDate")),
                    "survey_editdate": _xtext(row.get("EditDate")),
                    "survey_creator": _xtext(row.get("Creator")),
                    "survey_editor": _xtext(row.get("Editor")),
                    "survey_datos_originales": _row_json(row),
                    **datos_vinculacion_m01(conn, cedula),
                }
                # El primer GlobalID queda en la ficha; todos se preservan en referencias_survey123.
                if gid and (not existing or not _xtext(existing["survey_globalid"])):
                    values["survey_globalid"] = gid

                if existing:
                    update_values = dict(values)
                    update_values.pop("codigo_caso", None)
                    if additional_survey_reference:
                        # Otro GlobalID puede representar la misma ficha. Se conserva como
                        # referencia, sin reemplazar los metadatos principales del expediente.
                        for source_field in (
                            "survey_objectid", "survey_creationdate", "survey_editdate",
                            "survey_creator", "survey_editor", "survey_datos_originales",
                            "archivo_origen",
                        ):
                            update_values.pop(source_field, None)
                    if (
                        existing["id_persona_m01"] == update_values.get("id_persona_m01")
                        and existing["id_hogar_m01"] == update_values.get("id_hogar_m01")
                    ):
                        update_values["fecha_vinculacion_m01"] = existing["fecha_vinculacion_m01"]
                    if _changed(existing, update_values):
                        update_values["fecha_ultima_actualizacion"] = now_iso()
                        update_values["actualizado_por"] = user_name
                        _update(conn, "casos", "id_caso", existing["id_caso"], update_values)
                        counters["casos_actualizados"] += 1
                    else:
                        counters["casos_sin_cambios"] += 1
                    case_id = existing["id_caso"]
                else:
                    case_id = generate_id("CASO")
                    now = now_iso()
                    values.update({
                        "id_caso": case_id,
                        "confidencial": 0,
                        "fecha_creacion": now,
                        "fecha_ultima_actualizacion": now,
                        "actualizado_por": user_name,
                    })
                    columns = list(values.keys())
                    conn.execute(
                        f"INSERT INTO casos ({', '.join(columns)}) "
                        f"VALUES ({', '.join(['?'] * len(columns))})",
                        list(values.values()),
                    )
                    _assign_control_code(conn, "casos", "id_caso", case_id, "id_control_m8", "M8-C-")
                    register_state_change(
                        conn, case_id, None, values["estado_actual"], user_name, "Importación Survey123"
                    )
                    counters["casos_nuevos"] += 1

                _upsert_survey_reference(
                    conn, "caso", case_id, case_id, objectid, gid, "", form_number, source_file
                )
            except Exception as exc:
                counters["errores"] += 1
                errores.append(f"Caso fila {index + 2}: {exc}")

        for index, row in ubi_df.iterrows():
            try:
                gid = _xtext(row.get("GlobalID"))
                parent_gid = _xtext(row.get("ParentGlobalID"))
                objectid = _xtext(row.get("ObjectID"))
                form_number = _xtext(row.get("numFormularioUQ"))
                case = _parent_case(conn, parent_gid, form_number)
                if not case:
                    _save_pending_survey_record(
                        conn, "ubicacion", row, objectid, gid, parent_gid,
                        form_number, source_file,
                        "Caso padre no incluido o todavía no importado",
                    )
                    counters["pendientes_vinculacion"] += 1
                    continue
                values = {
                    "id_caso": case["id_caso"],
                    "survey_globalid": gid or None,
                    "survey_parentglobalid": parent_gid or None,
                    "survey_objectid": objectid,
                    "numero_formulario": form_number,
                    "provincia_codigo": _xcode(row.get("Provincia de la queja (cod.)"), 2),
                    "distrito_codigo": _xcode(row.get("Distrito de la queja (cod.)"), 4),
                    "corregimiento_codigo": _xcode(row.get("Corregimiento de la queja (cod.)"), 6),
                    "comunidad_codigo": _xcode(row.get("Comunidad de la queja (cod.)"), 9),
                    "provincia": _xtext(row.get("Provincia de la queja")),
                    "distrito": _xtext(row.get("Distrito de la queja")),
                    "corregimiento": _xtext(row.get("Corregimiento de la queja")),
                    "comunidad": _xtext(row.get("Comunidad de la queja")),
                    "direccion": _xtext(row.get("Dirección de la queja")),
                    "referencia_afectacion": _xtext(row.get("Referencia de la afectación")),
                    "otra_referencia": _xtext(row.get("Otra referencia de la afectación")),
                    "llegada_lugar": _xtext(row.get("Llegada al lugar de afectación")),
                    "coordenadas_geograficas": _xtext(row.get("Coordenadas geográficas del sitio")),
                    "coordenada_norte_utm": None if pd.isna(row.get("Coordenadas Norte UTM")) else row.get("Coordenadas Norte UTM"),
                    "coordenada_este_utm": None if pd.isna(row.get("Coordenadas Este UTM")) else row.get("Coordenadas Este UTM"),
                    "latitud": None if pd.isna(row.get("y")) else row.get("y"),
                    "longitud": None if pd.isna(row.get("x")) else row.get("x"),
                    "survey_creationdate": _xtext(row.get("CreationDate")),
                    "survey_editdate": _xtext(row.get("EditDate")),
                    "survey_creator": _xtext(row.get("Creator")),
                    "survey_editor": _xtext(row.get("Editor")),
                    "survey_datos_originales": _row_json(row),
                    "archivo_origen": source_file,
                    "fecha_importacion": now_iso(),
                }
                existing = _existing_source_record(
                    conn, "ubicaciones_survey123", gid, objectid, "numero_formulario", form_number
                )
                if existing:
                    values["fecha_importacion"] = existing["fecha_importacion"]
                    if _changed(existing, values):
                        _update(conn, "ubicaciones_survey123", "id_ubicacion", existing["id_ubicacion"], values)
                        counters["ubicaciones_actualizadas"] += 1
                    else:
                        counters["ubicaciones_sin_cambios"] += 1
                    location_id = existing["id_ubicacion"]
                else:
                    location_id = generate_id("UBI")
                    values["id_ubicacion"] = location_id
                    columns = list(values.keys())
                    conn.execute(
                        f"INSERT INTO ubicaciones_survey123 ({', '.join(columns)}) "
                        f"VALUES ({', '.join(['?'] * len(columns))})",
                        list(values.values()),
                    )
                    _assign_control_code(
                        conn, "ubicaciones_survey123", "id_ubicacion", location_id,
                        "id_control_ubicacion", "M8-U-"
                    )
                    counters["ubicaciones_nuevas"] += 1

                _update(conn, "casos", "id_caso", case["id_caso"], {
                    "provincia_hecho_codigo": values["provincia_codigo"],
                    "distrito_hecho_codigo": values["distrito_codigo"],
                    "corregimiento_hecho_codigo": values["corregimiento_codigo"],
                    "lugar_poblado_hecho_codigo": values["comunidad_codigo"],
                    "provincia_hecho": values["provincia"],
                    "distrito_hecho": values["distrito"],
                    "corregimiento_hecho": values["corregimiento"],
                    "lugar_poblado_hecho": values["comunidad"],
                    "direccion_hecho": values["direccion"],
                    "referencia_afectacion": values["referencia_afectacion"],
                    "otra_referencia_afectacion": values["otra_referencia"],
                    "llegada_lugar_afectacion": values["llegada_lugar"],
                    "tiene_coordenadas_hecho": values["coordenadas_geograficas"],
                    "coordenada_norte_utm": values["coordenada_norte_utm"],
                    "coordenada_este_utm": values["coordenada_este_utm"],
                    "latitud": values["latitud"],
                    "longitud": values["longitud"],
                })
                _upsert_survey_reference(
                    conn, "ubicacion", case["id_caso"], location_id,
                    objectid, gid, parent_gid, form_number, source_file
                )
            except Exception as exc:
                counters["errores"] += 1
                errores.append(f"Ubicación fila {index + 2}: {exc}")

        for index, row in seg_df.iterrows():
            try:
                gid = _xtext(row.get("GlobalID"))
                parent_gid = _xtext(row.get("ParentGlobalID"))
                objectid = _xtext(row.get("ObjectID"))
                form_number = _xtext(row.get("numFormularioS"))
                case = _parent_case(conn, parent_gid, form_number)
                if not case:
                    _save_pending_survey_record(
                        conn, "seguimiento", row, objectid, gid, parent_gid,
                        form_number, source_file,
                        "Caso padre no incluido o todavía no importado",
                    )
                    counters["pendientes_vinculacion"] += 1
                    continue
                values = {
                    "id_caso": case["id_caso"],
                    "fecha_actuacion": _xdate(row.get("Fecha")) or case["fecha_registro"],
                    "hora_actuacion": _xtime(row.get("Fecha")),
                    "tipo_actuacion": "Seguimiento importado",
                    "descripcion": _xtext(row.get("Descripción y seguimiento de la")) or "Sin descripción",
                    "resultado": "",
                    "responsable_ejecutor": _xtext(row.get("Creator")) or user_name,
                    "usuario_registro": user_name,
                    "estado_anterior": case["estado_actual"],
                    "estado_posterior": case["estado_actual"],
                    "proxima_accion": "",
                    "fecha_compromiso": None,
                    "estado_actividad": "Completada",
                    "visible_solicitante": 1,
                    "fecha_registro_sistema": now_iso(),
                    "clasificacion_survey": _xtext(row.get("clasificacionS")),
                    "numero_formulario_survey": form_number,
                    "survey_globalid": gid or None,
                    "survey_parentglobalid": parent_gid or None,
                    "survey_objectid": objectid,
                    "survey_creationdate": _xtext(row.get("CreationDate")),
                    "survey_editdate": _xtext(row.get("EditDate")),
                    "survey_creator": _xtext(row.get("Creator")),
                    "survey_editor": _xtext(row.get("Editor")),
                    "survey_datos_originales": _row_json(row),
                    "archivo_origen": source_file,
                }
                existing = _existing_source_record(
                    conn, "seguimientos", gid, objectid, "numero_formulario_survey", form_number
                )
                if existing:
                    values["fecha_registro_sistema"] = existing["fecha_registro_sistema"]
                    if _changed(existing, values):
                        _update(conn, "seguimientos", "id_seguimiento", existing["id_seguimiento"], values)
                        counters["seguimientos_actualizados"] += 1
                    else:
                        counters["seguimientos_sin_cambios"] += 1
                    followup_id = existing["id_seguimiento"]
                else:
                    followup_id = generate_id("SEG")
                    values["id_seguimiento"] = followup_id
                    columns = list(values.keys())
                    conn.execute(
                        f"INSERT INTO seguimientos ({', '.join(columns)}) "
                        f"VALUES ({', '.join(['?'] * len(columns))})",
                        list(values.values()),
                    )
                    _assign_control_code(
                        conn, "seguimientos", "id_seguimiento", followup_id,
                        "id_control_seguimiento", "M8-S-"
                    )
                    counters["seguimientos_nuevos"] += 1
                _upsert_survey_reference(
                    conn, "seguimiento", case["id_caso"], followup_id,
                    objectid, gid, parent_gid, form_number, source_file
                )
            except Exception as exc:
                counters["errores"] += 1
                errores.append(f"Seguimiento fila {index + 2}: {exc}")

        conn.execute(
            """
            INSERT INTO importaciones_survey123 (
                id_importacion, nombre_archivo, fecha_importacion, usuario,
                casos_nuevos, casos_actualizados, casos_sin_cambios,
                ubicaciones_nuevas, ubicaciones_actualizadas, ubicaciones_sin_cambios,
                seguimientos_nuevos, seguimientos_actualizados, seguimientos_sin_cambios,
                pendientes_vinculacion, errores, detalle_errores
            ) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
            """,
            (
                generate_id("IMP"), source_file, now_iso(), user_name,
                counters["casos_nuevos"], counters["casos_actualizados"], counters["casos_sin_cambios"],
                counters["ubicaciones_nuevas"], counters["ubicaciones_actualizadas"], counters["ubicaciones_sin_cambios"],
                counters["seguimientos_nuevos"], counters["seguimientos_actualizados"], counters["seguimientos_sin_cambios"],
                counters["pendientes_vinculacion"], counters["errores"], json.dumps(errores, ensure_ascii=False),
            ),
        )
    counters["detalle_errores"] = errores
    return counters


# =========================================================
# BASE DE DATOS
# =========================================================

def initialize_database() -> None:
    with db_connection() as conn:
        conn.executescript(
            """
            CREATE TABLE IF NOT EXISTS casos (
                id_caso TEXT PRIMARY KEY,
                codigo_caso TEXT UNIQUE NOT NULL,
                fecha_recepcion TEXT NOT NULL,
                hora_recepcion TEXT,
                fecha_registro TEXT NOT NULL,
                hora_registro TEXT,
                registrado_por TEXT NOT NULL,

                clasificacion TEXT NOT NULL CHECK (clasificacion IN ('Consulta', 'Queja')),
                estado_actual TEXT NOT NULL,
                situacion_atencion TEXT,
                dentro_alcance INTEGER NOT NULL DEFAULT 1,
                motivo_fuera_alcance TEXT,

                tipo_identificacion_solicitante TEXT NOT NULL,
                confidencial INTEGER NOT NULL DEFAULT 0,
                nombre_solicitante TEXT,
                sexo TEXT,
                cedula TEXT,
                telefono TEXT,
                celular TEXT,
                correo TEXT,

                provincia_contacto TEXT,
                distrito_contacto TEXT,
                corregimiento_contacto TEXT,
                lugar_poblado_contacto TEXT,
                direccion_contacto TEXT,

                medio_recepcion TEXT NOT NULL,
                otro_medio_recepcion TEXT,
                recibido_por TEXT NOT NULL,
                responsable_origen TEXT,
                tema TEXT,
                tipo_queja TEXT,
                descripcion TEXT NOT NULL,
                presentado_anteriormente INTEGER NOT NULL DEFAULT 0,
                referencia_caso_anterior TEXT,
                respuesta_inmediata TEXT,
                propuesta_solucion TEXT,

                provincia_hecho TEXT,
                distrito_hecho TEXT,
                corregimiento_hecho TEXT,
                lugar_poblado_hecho TEXT,
                direccion_hecho TEXT,
                latitud REAL,
                longitud REAL,

                supervisor TEXT,
                responsable_principal TEXT,
                responsable_apoyo_1 TEXT,
                responsable_apoyo_2 TEXT,
                fecha_asignacion TEXT,
                asignado_por TEXT,
                motivo_asignacion TEXT,

                fecha_acuse TEXT,
                medio_acuse TEXT,
                resultado_acuse TEXT,
                realizado_acuse_por TEXT,

                respuesta_final TEXT,
                acepta_respuesta TEXT,
                nivel_satisfaccion TEXT,
                comentario_insatisfaccion TEXT,
                comentarios_finales TEXT,

                recomendacion_cierre TEXT,
                fecha_recomendacion_cierre TEXT,
                recomendada_por TEXT,
                decision_supervisor TEXT,
                observacion_supervisor TEXT,
                fecha_decision_supervisor TEXT,

                fecha_comunicacion_cierre TEXT,
                medio_comunicacion_cierre TEXT,
                comunicado_por TEXT,
                resultado_comunicacion_cierre TEXT,

                fecha_cierre TEXT,
                cerrado_por TEXT,
                motivo_cierre TEXT,
                firma_conformidad INTEGER NOT NULL DEFAULT 0,
                negativa_firma INTEGER NOT NULL DEFAULT 0,
                testigo_cierre TEXT,

                fecha_creacion TEXT NOT NULL,
                fecha_ultima_actualizacion TEXT NOT NULL,
                actualizado_por TEXT NOT NULL
            );

            CREATE TABLE IF NOT EXISTS seguimientos (
                id_seguimiento TEXT PRIMARY KEY,
                id_caso TEXT NOT NULL,
                fecha_actuacion TEXT NOT NULL,
                hora_actuacion TEXT,
                tipo_actuacion TEXT NOT NULL,
                descripcion TEXT NOT NULL,
                resultado TEXT,
                responsable_ejecutor TEXT NOT NULL,
                usuario_registro TEXT NOT NULL,
                estado_anterior TEXT,
                estado_posterior TEXT,
                proxima_accion TEXT,
                fecha_compromiso TEXT,
                estado_actividad TEXT NOT NULL DEFAULT 'Pendiente',
                visible_solicitante INTEGER NOT NULL DEFAULT 0,
                fecha_registro_sistema TEXT NOT NULL,
                FOREIGN KEY (id_caso) REFERENCES casos(id_caso)
            );

            CREATE TABLE IF NOT EXISTS historial_estados (
                id_historial TEXT PRIMARY KEY,
                id_caso TEXT NOT NULL,
                estado_anterior TEXT,
                estado_nuevo TEXT NOT NULL,
                fecha_cambio TEXT NOT NULL,
                usuario_cambio TEXT NOT NULL,
                motivo TEXT,
                id_seguimiento_origen TEXT,
                FOREIGN KEY (id_caso) REFERENCES casos(id_caso),
                FOREIGN KEY (id_seguimiento_origen) REFERENCES seguimientos(id_seguimiento)
            );

            CREATE TABLE IF NOT EXISTS comunicaciones (
                id_comunicacion TEXT PRIMARY KEY,
                id_caso TEXT NOT NULL,
                tipo_comunicacion TEXT NOT NULL,
                fecha TEXT NOT NULL,
                hora TEXT,
                medio TEXT NOT NULL,
                destinatario TEXT,
                realizada_por TEXT NOT NULL,
                resultado_contacto TEXT NOT NULL,
                descripcion TEXT,
                es_acuse_recibo INTEGER NOT NULL DEFAULT 0,
                es_comunicacion_avance INTEGER NOT NULL DEFAULT 0,
                es_comunicacion_cierre INTEGER NOT NULL DEFAULT 0,
                fecha_registro_sistema TEXT NOT NULL,
                FOREIGN KEY (id_caso) REFERENCES casos(id_caso)
            );

            CREATE TABLE IF NOT EXISTS revisiones (
                id_revision TEXT PRIMARY KEY,
                id_caso TEXT NOT NULL,
                tipo_revision TEXT NOT NULL,
                fecha_solicitud TEXT NOT NULL,
                solicitada_por TEXT,
                motivo TEXT NOT NULL,
                revisor TEXT,
                fecha_revision TEXT,
                decision TEXT,
                observaciones TEXT,
                estado_revision TEXT NOT NULL DEFAULT 'Pendiente',
                FOREIGN KEY (id_caso) REFERENCES casos(id_caso)
            );

            CREATE TABLE IF NOT EXISTS documentos (
                id_documento TEXT PRIMARY KEY,
                id_caso TEXT NOT NULL,
                id_seguimiento TEXT,
                id_comunicacion TEXT,
                tipo_documento TEXT NOT NULL,
                nombre_archivo TEXT NOT NULL,
                ruta_archivo TEXT NOT NULL,
                fecha_carga TEXT NOT NULL,
                cargado_por TEXT NOT NULL,
                FOREIGN KEY (id_caso) REFERENCES casos(id_caso),
                FOREIGN KEY (id_seguimiento) REFERENCES seguimientos(id_seguimiento),
                FOREIGN KEY (id_comunicacion) REFERENCES comunicaciones(id_comunicacion)
            );

            CREATE TABLE IF NOT EXISTS auditoria (
                id_auditoria TEXT PRIMARY KEY,
                id_caso TEXT NOT NULL,
                accion TEXT NOT NULL,
                usuario TEXT NOT NULL,
                fecha_hora TEXT NOT NULL,
                datos_anteriores TEXT,
                datos_nuevos TEXT,
                detalle TEXT,
                FOREIGN KEY (id_caso) REFERENCES casos(id_caso)
            );

            CREATE INDEX IF NOT EXISTS idx_casos_codigo ON casos(codigo_caso);
            CREATE INDEX IF NOT EXISTS idx_casos_estado ON casos(estado_actual);
            CREATE INDEX IF NOT EXISTS idx_casos_responsable ON casos(responsable_principal);
            CREATE INDEX IF NOT EXISTS idx_seguimientos_caso ON seguimientos(id_caso);
            CREATE INDEX IF NOT EXISTS idx_historial_caso ON historial_estados(id_caso);
            """
        )


def ensure_column(table_name: str, column_name: str, column_definition: str) -> None:
    """Agrega una columna sin afectar bases existentes ni datos ya registrados."""
    with db_connection() as conn:
        existing = {row["name"] for row in conn.execute(f"PRAGMA table_info({table_name})").fetchall()}
        if column_name not in existing:
            conn.execute(
                f"ALTER TABLE {table_name} ADD COLUMN {column_name} {column_definition}"
            )


initialize_database()
ensure_column("casos", "hora_registro", "TEXT")
ensure_column("casos", "responsable_origen", "TEXT")
ensure_column("casos", "origen_gestion", "TEXT")
ensure_column("casos", "cargo_registra", "TEXT")
ensure_column("casos", "respecto_a", "TEXT")
ensure_column("casos", "respecto_otro", "TEXT")
ensure_column("casos", "atendida_por", "TEXT")
ensure_column("casos", "resultado_atencion", "TEXT")
ensure_column("casos", "remitida_oficina", "TEXT")
ensure_column("casos", "requiere_asesoria_juridica", "INTEGER NOT NULL DEFAULT 0")
ensure_column("casos", "fecha_remision_juridica", "TEXT")
ensure_column("casos", "estado_juridico", "TEXT")
ensure_column("casos", "cierre_juridico_mayor_6_meses", "INTEGER NOT NULL DEFAULT 0")
ensure_column("casos", "visto_bueno_supervisor", "TEXT")
ensure_column("casos", "fecha_visto_bueno", "TEXT")
ensure_column("casos", "observacion_visto_bueno", "TEXT")
ensure_column("casos", "acp_participa_cierre_contratista", "TEXT")
ensure_column("casos", "id_persona_m01", "TEXT")
ensure_column("casos", "id_hogar_m01", "TEXT")
ensure_column("casos", "pertenece_proyecto", "INTEGER NOT NULL DEFAULT 0")
ensure_column("casos", "estado_vinculacion_m01", "TEXT")
ensure_column("casos", "fecha_vinculacion_m01", "TEXT")
with db_connection() as conn:
    ensure_import_schema(conn)
    # Migración no destructiva: los valores históricos siguen disponibles y
    # se copian a la nueva denominación general de tipo de caso.
    conn.execute(
        """
        UPDATE casos
        SET tipo_caso = COALESCE(NULLIF(tipo_caso, ''), tipo_queja),
            tipo_caso_otro = COALESCE(NULLIF(tipo_caso_otro, ''), tipo_queja_otro)
        WHERE COALESCE(tipo_caso, '') = '' OR COALESCE(tipo_caso_otro, '') = ''
        """
    )


# =========================================================
# INTEGRACIÓN SIMULADA CON M01 · PERSONAS Y HOGARES
# =========================================================

def normalizar_cedula(value: Any) -> str:
    """Normaliza la cédula para comparar sin guiones, espacios o puntos."""
    return re.sub(r"[^0-9A-Za-z]", "", str(value or "")).upper()


def initialize_m01_simulation() -> None:
    """Crea una estructura mínima compatible con M01 sin alterar sus datos futuros."""
    with db_connection() as conn:
        conn.executescript(
            """
            CREATE TABLE IF NOT EXISTS m01_hogares (
                id_hogar TEXT PRIMARY KEY,
                codigo_hogar_campo TEXT,
                nombre_referencia_hogar TEXT,
                id_lugar_poblado TEXT,
                zona TEXT,
                tipo_afectacion TEXT,
                tipo_desplazamiento TEXT,
                nivel_prioridad_social TEXT
            );

            CREATE TABLE IF NOT EXISTS m01_personas (
                id_persona TEXT PRIMARY KEY,
                id_hogar TEXT NOT NULL,
                nombres TEXT,
                apellidos TEXT,
                documento_identidad TEXT,
                documento_normalizado TEXT,
                telefono TEXT,
                sexo TEXT,
                fecha_nacimiento TEXT,
                parentesco TEXT,
                jefe_hogar INTEGER NOT NULL DEFAULT 0,
                vive_en_hogar INTEGER NOT NULL DEFAULT 1,
                FOREIGN KEY (id_hogar) REFERENCES m01_hogares(id_hogar)
            );

            CREATE UNIQUE INDEX IF NOT EXISTS idx_m01_persona_documento
            ON m01_personas(documento_normalizado)
            WHERE documento_normalizado IS NOT NULL AND documento_normalizado <> '';
            """
        )

        total = conn.execute("SELECT COUNT(*) AS total FROM m01_hogares").fetchone()["total"]
        if ENABLE_DEMO_DATA and total == 0:
            hogares = [
                ("HOG-0001", "PA-RI-001", "Hogar Anicasio Carrión", "COM-0001", "Río Indio", "Físico", "Físico-económico", "Alta"),
                ("HOG-0002", "PA-RI-002", "Hogar María González", "COM-0002", "Río Indio", "Económico", "Económico", "Media"),
                ("HOG-0003", "PA-CH-003", "Hogar José Pérez", "COM-0003", "Cuenca Oeste", "Físico", "Físico", "Alta"),
                ("HOG-0004", "PA-CH-004", "Hogar Elena Rodríguez", "COM-0004", "Cuenca Oeste", "Económico", "Económico", "Baja"),
            ]
            conn.executemany(
                """
                INSERT INTO m01_hogares (
                    id_hogar, codigo_hogar_campo, nombre_referencia_hogar,
                    id_lugar_poblado, zona, tipo_afectacion,
                    tipo_desplazamiento, nivel_prioridad_social
                ) VALUES (?, ?, ?, ?, ?, ?, ?, ?)
                """,
                hogares,
            )

            personas = [
                ("PER-0001", "HOG-0001", "Anicasio", "Carrión", "3-88-124", normalizar_cedula("3-88-124"), "6000-0101", "Masculino", "1970-03-14", "Jefe de hogar", 1, 1),
                ("PER-0002", "HOG-0002", "María", "González", "8-765-432", normalizar_cedula("8-765-432"), "6000-0202", "Femenino", "1982-07-20", "Jefa de hogar", 1, 1),
                ("PER-0003", "HOG-0003", "José", "Pérez", "4-123-987", normalizar_cedula("4-123-987"), "6000-0303", "Masculino", "1976-11-02", "Jefe de hogar", 1, 1),
                ("PER-0004", "HOG-0004", "Elena", "Rodríguez", "8-999-111", normalizar_cedula("8-999-111"), "6000-0404", "Femenino", "1990-01-15", "Jefa de hogar", 1, 1),
                ("PER-0005", "HOG-0001", "Luis", "Carrión", "3-88-125", normalizar_cedula("3-88-125"), "6000-0102", "Masculino", "1995-09-09", "Hijo", 0, 1),
            ]
            conn.executemany(
                """
                INSERT INTO m01_personas (
                    id_persona, id_hogar, nombres, apellidos,
                    documento_identidad, documento_normalizado,
                    telefono, sexo, fecha_nacimiento, parentesco,
                    jefe_hogar, vive_en_hogar
                ) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
                """,
                personas,
            )


def buscar_persona_m01_por_cedula(
    conn: sqlite3.Connection,
    cedula: str,
) -> Optional[sqlite3.Row]:
    cedula_normalizada = normalizar_cedula(cedula)
    if not cedula_normalizada:
        return None
    return conn.execute(
        """
        SELECT
            p.id_persona,
            p.id_hogar,
            p.nombres,
            p.apellidos,
            p.documento_identidad,
            h.codigo_hogar_campo,
            h.nombre_referencia_hogar,
            h.zona,
            h.tipo_afectacion,
            h.tipo_desplazamiento,
            h.nivel_prioridad_social
        FROM m01_personas p
        JOIN m01_hogares h ON h.id_hogar = p.id_hogar
        WHERE p.documento_normalizado = ?
        """,
        (cedula_normalizada,),
    ).fetchone()


def datos_vinculacion_m01(
    conn: sqlite3.Connection,
    cedula: str,
) -> Dict[str, Any]:
    persona = buscar_persona_m01_por_cedula(conn, cedula)
    if persona:
        return {
            "id_persona_m01": persona["id_persona"],
            "id_hogar_m01": persona["id_hogar"],
            "pertenece_proyecto": 1,
            "estado_vinculacion_m01": "Vinculado con persona y hogar de M01",
            "fecha_vinculacion_m01": now_iso(),
        }
    return {
        "id_persona_m01": None,
        "id_hogar_m01": None,
        "pertenece_proyecto": 0,
        "estado_vinculacion_m01": (
            "Persona externa al proyecto o cédula no encontrada en M01"
            if normalize_text(cedula)
            else "No evaluado: caso sin cédula"
        ),
        "fecha_vinculacion_m01": now_iso(),
    }


def vincular_caso_m01(
    conn: sqlite3.Connection,
    case_id: str,
    cedula: str,
) -> Dict[str, Any]:
    relation = datos_vinculacion_m01(conn, cedula)
    _update(conn, "casos", "id_caso", case_id, relation)
    return relation


def seed_sample_cases_if_empty() -> None:
    """Crea casos de prueba solo cuando la tabla casos está vacía."""
    with db_connection() as conn:
        total = conn.execute("SELECT COUNT(*) AS total FROM casos").fetchone()["total"]
        if total > 0:
            return

        samples = [
            {
                "codigo_caso": "CQ-2026-00001",
                "cedula": "3-88-124",
                "nombre": "Anicasio Carrión",
                "clasificacion": "Queja",
                "tema": "Afectación de cultivos",
                "descripcion": "Daños reportados en cultivos durante actividades de campo del proyecto.",
                "estado": "En atención",
            },
            {
                "codigo_caso": "CQ-2026-00002",
                "cedula": "8-765-432",
                "nombre": "María González",
                "clasificacion": "Consulta",
                "tema": "Información sobre cronograma",
                "descripcion": "Solicita información sobre próximas actividades del programa en su comunidad.",
                "estado": "Registrada",
            },
            {
                "codigo_caso": "CQ-2026-00003",
                "cedula": "9-999-999",
                "nombre": "Carlos Méndez",
                "clasificacion": "Queja",
                "tema": "Tránsito de vehículos",
                "descripcion": "Reporta molestias por tránsito de vehículos fuera de la huella directa del proyecto.",
                "estado": "En atención",
            },
            {
                "codigo_caso": "CQ-2026-00004",
                "cedula": "",
                "nombre": "Rosa Martínez",
                "clasificacion": "Consulta",
                "tema": "Acceso a información",
                "descripcion": "Consulta de una persona externa sobre acceso a información pública del programa.",
                "estado": "Registrada",
            },
        ]

        for sample in samples:
            case_id = generate_id("CASO")
            relation = datos_vinculacion_m01(conn, sample["cedula"])
            now = now_iso()
            conn.execute(
                """
                INSERT INTO casos (
                    id_caso, codigo_caso, fecha_recepcion, hora_recepcion,
                    fecha_registro, hora_registro, registrado_por,
                    clasificacion, estado_actual, situacion_atencion,
                    dentro_alcance, motivo_fuera_alcance,
                    tipo_identificacion_solicitante, confidencial,
                    nombre_solicitante, sexo, cedula, telefono, celular, correo,
                    provincia_contacto, distrito_contacto, corregimiento_contacto,
                    lugar_poblado_contacto, direccion_contacto,
                    medio_recepcion, otro_medio_recepcion, recibido_por,
                    responsable_origen, tema, tipo_queja, descripcion,
                    presentado_anteriormente, referencia_caso_anterior,
                    respuesta_inmediata, propuesta_solucion,
                    provincia_hecho, distrito_hecho, corregimiento_hecho,
                    lugar_poblado_hecho, direccion_hecho,
                    id_persona_m01, id_hogar_m01, pertenece_proyecto,
                    estado_vinculacion_m01, fecha_vinculacion_m01,
                    fecha_creacion, fecha_ultima_actualizacion, actualizado_por
                ) VALUES (
                    ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?,
                    ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?,
                    ?, ?, ?, ?, ?, ?, ?, ?
                )
                """,
                (
                    case_id, sample["codigo_caso"], "2026-06-10", "08:30",
                    "2026-06-10", "08:30", "Usuario demostración",
                    sample["clasificacion"], sample["estado"], "",
                    1, "",
                    "Identificado", 0,
                    sample["nombre"], "", sample["cedula"], "", "", "",
                    "COLÓN", "CHAGRES", "LA ENCANTADA",
                    "EL TORNITO", "",
                    "Personalmente", "", "Equipo social",
                    "", sample["tema"], "", sample["descripcion"],
                    0, "", "Se informó que el caso sería registrado y atendido.", "",
                    "COLÓN", "CHAGRES", "LA ENCANTADA",
                    "EL TORNITO", "",
                    relation["id_persona_m01"], relation["id_hogar_m01"],
                    relation["pertenece_proyecto"],
                    relation["estado_vinculacion_m01"],
                    relation["fecha_vinculacion_m01"],
                    now, now, "Usuario demostración",
                ),
            )
            register_state_change(
                conn,
                case_id,
                None,
                sample["estado"],
                "Usuario demostración",
                "Caso simulado para validar la integración con M01.",
            )


initialize_m01_simulation()
if ENABLE_DEMO_DATA:
    seed_sample_cases_if_empty()


# =========================================================
# CATÁLOGOS DEL FORMULARIO OFICIAL
# =========================================================

# Los valores se tomaron de la hoja "choices" del XLSForm suministrado.
MEDIOS_RECEPCION = [
    "Nota",
    "Teléfono",
    "Correo electrónico",
    "Conversación con especialista",
    "Visita ORC",
    "Otro",
]

CONTACTO_PERSONA_OPTIONS = ["Teléfono", "Correo electrónico", "Personalmente"]
SEXO_OPTIONS = ["Hombre", "Mujer", "No identificado"]
CLASIFICACION_OPTIONS = ["Consulta", "Queja"]
SI_NO_OPTIONS = ["Sí", "No"]
ESTATUS_OPTIONS = ["Abierta", "Cerrada"]
TELEFONO_CELULAR_OPTIONS = ["Teléfono", "Celular", "Ambos"]

TEMAS_CASO = [
    "Avalúos",
    "Catrastro y titulación de tierras",
    "Conocer Oficina de Relación Comunitaria",
    "Estudios ambientales",
    "Oferta de predios",
    "Oferta de servicios",
    "Proyecto Lago de Río Indio",
    "Reasentamiento",
    "Tramite No Objeción",
    "UTEP, proyectos y otros temas",
]

TIPOS_CASO = [
    "Proyectos o construcciones",
    "Programas",
    "Estudios o actividad",
    "Otro",
]

TIPOS_CASO_QUEJA = TIPOS_CASO  # alias técnico de compatibilidad

REFERENCIAS_AFECTACION = ["Tienda", "Colegio", "Finca", "Otra referencia"]
CANTIDAD_RESPONSABLES_OPTIONS = [1, 2, 3]

OFICINAS_RECEPCION = [
    "BOCA DE RÍO INDIO",
    "EL CONGO",
    "EL LIMÓN",
    "LAS MARÍAS",
    "RÍO INDIO DE LOS CHORROS",
    "SAN CRISTÓBAL",
    "TRES HERMANAS",
    "PH",
    "ACP",
    "HI",
    "UTA",
    "PH/CONTRATISTA",
    "OTRO",
]

EQUIPOS_PH = [
    "Equipo de Manejo de Cuenca (PH-SM)",
    "Equipo de Seguimiento Ambiental de Proyectos (PH-SA)",
    "Equipo de Seguimiento de Reasentamiento (PH-SR)",
    "Equipo de Gestión de Tierras (PH-ST)",
]

SATISFACCION_OPTIONS = [
    "1 - Nada satisfecho",
    "2 - Poco satisfecho",
    "3 - Neutral / Indiferente",
    "4 - Satisfecho",
    "5 - Muy satisfecho",
]

ACEPTACION_RESPUESTA_OPTIONS = ["Pendiente", "Sí", "No"]

PROVINCIAS = [
    ("02", "COCLÉ"),
    ("03", "COLÓN"),
    ("13", "PANAMÁ OESTE"),
    ("08", "PANAMÁ"),
]

DISTRITOS = [
    ("0202", "ANTÓN", "02"),
    ("0203", "LA PINTADA", "02"),
    ("0206", "PENONOMÉ", "02"),
    ("0301", "COLÓN", "03"),
    ("0302", "CHAGRES", "03"),
    ("0303", "DONOSO", "03"),
    ("0304", "PORTOBELO", "03"),
    ("0306", "OMAR TORRIJOS HERRERA", "03"),
    ("0808", "PANAMÁ", "08"),
    ("1301", "ARRAIJÁN", "13"),
    ("1303", "CAPIRA", "13"),
    ("1307", "LA CHORRERA", "13"),
]

CORREGIMIENTOS = [
    ("020611", "BOCA DE TUCUU", "0206"),
    ("020612", "CANDELARIO OVALLE", "0206"),
    ("020604", "CHIGUIRÍ ARRIBA", "0206"),
    ("020302", "EL HARINO", "0203"),
    ("020205", "EL VALLE", "0202"),
    ("020613", "GENERAL VICTORIANO LORENZO", "0206"),
    ("020307", "LLANO NORTE", "0203"),
    ("020606", "PAJONAL", "0206"),
    ("020305", "PIEDRAS GORDAS", "0203"),
    ("020608", "RÍO INDIO", "0206"),
    ("020615", "RIECITO", "0206"),
    ("020616", "SAN MIGUEL", "0206"),
    ("020609", "TOABRÉ", "0206"),
    ("020610", "TULÚ", "0206"),
    ("030103", "BUENA VISTA", "0301"),
    ("030105", "CIRICITO", "0301"),
    ("030302", "COCLÉ DEL NORTE", "0303"),
    ("030106", "CRISTÓBAL", "0301"),
    ("030303", "EL GUÁSIMO", "0303"),
    ("030107", "ESCOBAL", "0301"),
    ("030304", "GOBEA", "0303"),
    ("030204", "LA ENCANTADA", "0302"),
    ("030108", "LIMÓN", "0301"),
    ("030405", "MARÍA CHIQUITA", "0304"),
    ("030301", "MIGUEL DE LA BORDA", "0303"),
    ("030602", "NUEVA ESPERANZA", "0306"),
    ("030109", "NUEVA PROVIDENCIA", "0301"),
    ("030110", "PUERTO PILÓN", "0301"),
    ("030305", "RÍO INDIO", "0303"),
    ("030111", "SABANITAS", "0301"),
    ("030112", "SALAMANCA", "0301"),
    ("030207", "SALUD", "0302"),
    ("030601", "SAN JOSÉ DEL GENERAL", "0306"),
    ("030113", "SAN JUAN", "0301"),
    ("030603", "SAN JUAN DE TURBE", "0306"),
    ("030114", "SANTA ROSA", "0301"),
    ("080822", "ALCALDE DÍAZ", "0808"),
    ("080814", "ANCÓN", "0808"),
    ("080824", "CAIMITILLO", "0808"),
    ("080815", "CHILIBRE", "0808"),
    ("080816", "LAS CUMBRES", "0808"),
    ("130703", "AMADOR", "1307"),
    ("130704", "AROSEMENA", "1307"),
    ("130101", "ARRAIJÁN", "1301"),
    ("130107", "BURUNGA", "1301"),
    ("130302", "CAIMITO", "1303"),
    ("130305", "CIRÍ DE LOS SOTOS", "1303"),
    ("130306", "CIRÍ GRANDE", "1303"),
    ("130705", "EL ARADO", "1307"),
    ("130307", "EL CACAO", "1303"),
    ("130709", "HERRERA", "1307"),
    ("130710", "HURTADO", "1307"),
    ("130711", "ITURRALDE", "1307"),
    ("130712", "LA REPRESA", "1307"),
    ("130308", "LA TRINIDAD", "1303"),
    ("130714", "MENDOZA", "1307"),
    ("130103", "NUEVO EMPERADOR", "1301"),
    ("130715", "OBALDÍA", "1307"),
    ("130104", "SANTA CLARA", "1301"),
    ("130313", "SANTA ROSA", "1303"),
]

CASE_STATES = [
    "Recibida",
    "Registrada",
    "Pendiente de asignación",
    "Asignada",
    "En atención",
    "Pendiente de aprobación",
    "Devuelta para atención",
    "Cierre aprobado",
    "Pendiente de comunicación",
    "Cerrada",
    "En revisión",
    "Fuera de alcance",
]

ATTENTION_SITUATIONS = [
    "",
    "En investigación",
    "Pendiente de información del solicitante",
    "Pendiente de tercero",
    "En inspección",
    "En evaluación técnica",
    "En asesoría jurídica",
    "Preparando respuesta",
]

FOLLOWUP_TYPES = [
    "Llamada",
    "Correo electrónico",
    "Reunión",
    "Inspección",
    "Visita de campo",
    "Revisión de antecedentes",
    "Solicitud de información",
    "Coordinación con tercero",
    "Evaluación técnica",
    "Preparación de respuesta",
    "Recomendación",
    "Remisión jurídica",
    "Otro",
]

RECEPTION_CHANNELS = MEDIOS_RECEPCION
COMMUNICATION_TYPES = [
    "Notificación de recepción",
    "Información de avance",
    "Solicitud de información",
    "Comunicación de resultado",
]


def _options_with_current(options: List[Any], current: Any, include_blank: bool = True) -> List[Any]:
    result: List[Any] = [""] if include_blank else []
    for option in options:
        if option not in result:
            result.append(option)
    if current not in (None, "") and current not in result:
        result.append(current)
    return result


def _option_index(options: List[Any], current: Any, default: int = 0) -> int:
    return options.index(current) if current in options else default


def _location_option(code: Any, label: Any) -> str:
    code_text = normalize_text(str(code or ""))
    label_text = normalize_text(str(label or ""))
    if code_text and label_text:
        return f"{code_text} · {label_text}"
    return label_text or code_text


def _decode_location_option(option: str) -> tuple[str, str]:
    option = normalize_text(option)
    if not option:
        return "", ""
    if " · " in option:
        code, label = option.split(" · ", 1)
        return code.strip(), label.strip()
    return "", option


def _location_options(items: List[tuple], current_code: Any, current_label: Any) -> List[str]:
    options = [""] + [_location_option(item[0], item[1]) for item in items]
    current = _location_option(current_code, current_label)
    if current and current not in options:
        options.append(current)
    return options


def _parse_multiselect(value: Any, available: List[str]) -> tuple[List[str], List[str]]:
    text_value = normalize_text(str(value or ""))
    selected: List[str] = []
    if text_value:
        selected = [part.strip() for part in re.split(r"\s*[;,|]\s*", text_value) if part.strip()]
    options = list(available)
    for item in selected:
        if item not in options:
            options.append(item)
    return options, selected

# =========================================================
# ESTILO E INTERFAZ CORPORATIVA
# =========================================================

COLOR_PRIMARIO_SOCIONAUT = "#073B5A"
COLOR_SECUNDARIO_SOCIONAUT = "#00A6A6"
COLOR_CORAL = "#F05A43"
COLOR_GRIS_CLARO = "#F4F7F9"
COLOR_BORDE = "#D6DEE6"

PAGE_HELP = {
    "Panel general": (
        "Consulta indicadores y los casos registrados recientemente."
    ),
    "Carga de información": (
        "Carga exclusivamente el archivo exportado desde Survey123. Los identificadores "
        "internos se generan automáticamente y las referencias de intercambio se conservan."
    ),
    "Registro de casos": (
        "Consulta los casos registrados, agrega un caso manualmente o actualiza su ficha."
    ),
    "Seguimiento": (
        "Registra actuaciones, resultados, compromisos e información comunicada a la persona "
        "dentro del seguimiento del caso."
    ),
    "Cierre": (
        "Registra la respuesta, la propuesta de cierre, el visto bueno del supervisor y la "
        "constancia final del caso."
    ),
    "Histórico": (
        "Consulta al final del proceso los casos, seguimientos, cierres y la trazabilidad completa."
    ),
    "Trazabilidad": (
        "Consulta la línea de tiempo construida con estados, seguimientos, documentos y auditoría."
    ),
}



def aplicar_estilos():
    st.markdown(
        f"""
        <style>
            :root {{
                --sir-primary: var(--primary-color, {COLOR_PRIMARIO_SOCIONAUT});
                --sir-accent: {COLOR_SECUNDARIO_SOCIONAUT};
                --sir-coral: {COLOR_CORAL};
                --sir-card: var(--secondary-background-color);
                --sir-text: var(--text-color);
                --sir-border: rgba(128,128,128,.28);
                --sir-shadow: rgba(0,0,0,.11);
            }}

            .block-container {{
                padding-top: 1.15rem;
                padding-bottom: 2.4rem;
                max-width: 1500px;
            }}

            .main-title {{
                font-size: clamp(1.5rem, 2.6vw, 2.25rem);
                font-weight: 950;
                color: var(--sir-primary);
                letter-spacing: -0.035em;
                margin-bottom: .15rem;
            }}

            .sub-title {{
                opacity: .76;
                margin-bottom: 1rem;
                font-size: .98rem;
            }}

            .section-card, .case-card {{
                background: var(--sir-card);
                color: var(--sir-text);
                border: 1px solid var(--sir-border);
                border-radius: 22px;
                box-shadow: 0 10px 28px var(--sir-shadow);
                padding: 1.05rem 1.2rem;
                margin-bottom: 1rem;
            }}

            .screen-help {{
                border-left: 5px solid var(--sir-accent);
                background: color-mix(in srgb, var(--sir-card) 84%, var(--sir-accent) 10%);
                border-radius: 16px;
                padding: .85rem 1rem;
                margin: .45rem 0 1rem 0;
                line-height: 1.45;
            }}

            .case-hero {{
                display:flex;
                justify-content:space-between;
                gap:1rem;
                align-items:flex-start;
            }}

            .case-kicker {{
                color:var(--sir-accent);
                font-weight:900;
                text-transform:uppercase;
                letter-spacing:.08em;
                font-size:.72rem;
            }}

            .case-title {{
                font-size:clamp(1.1rem,2vw,1.55rem);
                font-weight:950;
                letter-spacing:-.03em;
                margin:.15rem 0;
            }}

            .case-subtitle {{
                opacity:.72;
                font-size:.9rem;
            }}

            .chip {{
                display:inline-block;
                padding:.26rem .68rem;
                border-radius:999px;
                font-size:.79rem;
                font-weight:850;
                border:1px solid var(--sir-border);
                margin-left:.3rem;
                margin-bottom:.3rem;
                background: color-mix(in srgb, var(--sir-card) 80%, var(--sir-primary) 10%);
                color:var(--sir-text);
            }}

            .chip-accent {{
                background:rgba(0,166,166,.14);
                border-color:rgba(0,166,166,.38);
            }}

            .chip-coral {{
                background:rgba(240,90,67,.14);
                border-color:rgba(240,90,67,.38);
            }}

            div[data-testid="stMetric"] {{
                background:var(--sir-card);
                border:1px solid var(--sir-border);
                border-radius:18px;
                padding:1rem;
                box-shadow:0 8px 20px var(--sir-shadow);
            }}

            div[data-testid="stForm"] {{
                background:var(--sir-card);
                border:1px solid var(--sir-border);
                border-radius:22px;
                padding:1.15rem 1.25rem;
                box-shadow:0 10px 28px var(--sir-shadow);
            }}

            div[data-testid="stExpander"] {{
                border:1px solid var(--sir-border);
                border-radius:16px;
                overflow:hidden;
            }}

            .stButton > button, .stDownloadButton > button {{
                min-height:2.65rem;
                border-radius:14px !important;
                font-weight:800 !important;
                border:1px solid var(--sir-border) !important;
                transition:all 160ms ease-in-out;
                box-shadow:0 6px 16px rgba(0,0,0,.09);
            }}

            .stButton > button:hover, .stDownloadButton > button:hover {{
                transform:translateY(-1px);
                box-shadow:0 10px 22px rgba(0,0,0,.15);
            }}

            [data-testid="stSidebar"] {{
                border-right:1px solid var(--sir-border);
            }}

            [data-testid="stSidebar"] .stRadio > div {{
                gap:.25rem;
            }}

            [data-testid="stSidebar"] label[data-baseweb="radio"] {{
                border-radius:12px;
                padding:.38rem .55rem;
            }}

            .stTextInput label, .stSelectbox label, .stDateInput label,
            .stTimeInput label, .stNumberInput label, .stCheckbox label,
            .stTextArea label, .stRadio label, .stMultiSelect label,
            .stFileUploader label {{
                color:var(--sir-text) !important;
                font-weight:700;
            }}

            @media (max-width:768px) {{
                .case-hero {{ flex-direction:column; }}
                .section-card, .case-card {{ padding:.9rem; border-radius:18px; }}
                div[data-testid="stForm"] {{ padding:.9rem; border-radius:18px; }}
            }}
        </style>
        """,
        unsafe_allow_html=True,
    )


def mostrar_encabezado():
    st.markdown(
        '<div class="main-title">Módulo 8 · Casos</div>',
        unsafe_allow_html=True,
    )
    st.markdown(
        '<div class="sub-title">Sistema de Información para Reasentamiento · ACP · '
        'Registro, clasificación, seguimiento, aprobación y cierre</div>',
        unsafe_allow_html=True,
    )


def mostrar_ayuda_pantalla(nombre_pantalla: str) -> None:
    texto = PAGE_HELP.get(nombre_pantalla, "")
    if texto:
        st.markdown(
            f"<div class='screen-help'>💡 {texto}</div>",
            unsafe_allow_html=True,
        )


def mostrar_resumen_caso(case: sqlite3.Row) -> None:
    codigo = case["codigo_caso"] or "Sin código"
    clasificacion = case["clasificacion"] or "Sin clasificación"
    estado = case["estado_actual"] or "Sin estado"
    responsable = case["responsable_principal"] or "Sin responsable asignado"
    vinculo = (
        f"Proyecto · {case['id_persona_m01']} · {case['id_hogar_m01']}"
        if case["pertenece_proyecto"]
        else "Persona externa al M01 / sin coincidencia"
    )
    descripcion = (case["descripcion"] or "").strip()
    if len(descripcion) > 180:
        descripcion = descripcion[:177] + "..."

    st.markdown(
        f"""
        <div class="case-card">
            <div class="case-hero">
                <div>
                    <div class="case-kicker">Expediente seleccionado</div>
                    <div class="case-title">{codigo}</div>
                    <div class="case-subtitle">{descripcion or "Sin descripción registrada."}</div>
                </div>
                <div>
                    <span class="chip chip-accent">{clasificacion}</span>
                    <span class="chip chip-coral">{estado}</span>
                    <span class="chip">{responsable}</span>
                    <span class="chip chip-accent">{vinculo}</span>
                </div>
            </div>
        </div>
        """,
        unsafe_allow_html=True,
    )


aplicar_estilos()
mostrar_encabezado()


# =========================================================
# SESIÓN Y NAVEGACIÓN
# =========================================================

if "current_user" not in st.session_state:
    st.session_state.current_user = "Usuario SIR"

with st.sidebar:
    st.title("Módulo 8 · Casos")
    st.caption("Controles del módulo")
    page = st.radio(
        "Pantalla de trabajo",
        [
            "Panel general",
            "Carga de información",
            "Registro de casos",
            "Seguimiento",
            "Cierre",
            "Histórico",
        ],
        help="Selecciona el proceso que deseas realizar.",
    )

    st.markdown("---")
    st.subheader("Sesión")
    st.session_state.current_user = st.text_input(
        "Usuario activo",
        value=st.session_state.current_user,
        help="Este nombre se utilizará en la auditoría y en los registros de actividad.",
    )

    st.markdown("---")
    st.caption(f"Base de datos: {DB_PATH.name}")
    st.caption("Los identificadores y registros técnicos se conservan automáticamente.")


def case_selector(label: str = "Seleccione un caso") -> Optional[str]:
    with db_connection() as conn:
        rows = conn.execute(
            """
            SELECT id_caso, codigo_caso, clasificacion, estado_actual, descripcion
            FROM casos
            ORDER BY fecha_creacion DESC
            """
        ).fetchall()

    if not rows:
        st.info("Todavía no hay casos registrados.")
        return None

    options = {
        f"{r['codigo_caso']} · {r['clasificacion']} · {r['estado_actual']} · {r['descripcion'][:60]}": r["id_caso"]
        for r in rows
    }
    selected_label = st.selectbox(
        label,
        list(options.keys()),
        help="Busca el código, clasificación, estado o descripción del caso.",
    )
    selected_id = options[selected_label]
    with db_connection() as conn:
        selected_case = get_case(conn, selected_id)
    if selected_case:
        mostrar_resumen_caso(selected_case)
    return selected_id



# =========================================================
# CONSULTA, HISTÓRICOS Y EXPORTACIONES
# =========================================================

TABLE_EXPORT_ORDER = [
    "casos",
    "seguimientos",
    "comunicaciones",
    "historial_estados",
    "revisiones",
    "documentos",
    "auditoria",
    "ubicaciones_survey123",
    "importaciones_survey123",
    "referencias_survey123",
    "registros_survey_pendientes",
    "m01_hogares",
    "m01_personas",
]


def sql_table_exists(conn: sqlite3.Connection, table_name: str) -> bool:
    return conn.execute(
        "SELECT 1 FROM sqlite_master WHERE type='table' AND name = ?",
        (table_name,),
    ).fetchone() is not None


def read_table_df(table_name: str) -> pd.DataFrame:
    with db_connection() as conn:
        if not sql_table_exists(conn, table_name):
            return pd.DataFrame()
        return pd.read_sql_query(f'SELECT * FROM "{table_name}"', conn)


def read_query_df(query: str, params: tuple = ()) -> pd.DataFrame:
    with db_connection() as conn:
        return pd.read_sql_query(query, conn, params=params)


def get_case_options() -> Dict[str, str]:
    with db_connection() as conn:
        rows = conn.execute(
            """
            SELECT id_caso, id_control_m8, codigo_caso, clasificacion, estado_actual,
                   COALESCE(nombre_solicitante, '') AS nombre_solicitante
            FROM casos
            ORDER BY fecha_registro DESC, codigo_caso DESC
            """
        ).fetchall()
    return {
        (
            f"{row['codigo_caso']} · {row['clasificacion']} · "
            f"{row['estado_actual']} · {row['nombre_solicitante'] or 'Sin nombre registrado'}"
        ): row["id_caso"]
        for row in rows
    }


def select_case_id(label: str, key: str) -> Optional[str]:
    options = get_case_options()
    if not options:
        st.warning("No hay casos registrados.")
        return None
    selected = st.selectbox(label, list(options.keys()), key=key)
    case_id = options[selected]
    with db_connection() as conn:
        case = get_case(conn, case_id)
    if case:
        mostrar_resumen_caso(case)
    return case_id


def render_work_mode(options: List[str], key: str) -> str:
    return st.radio(
        "Sección de trabajo",
        options,
        horizontal=True,
        key=key,
        help="Cambia entre el histórico, la captura de nuevos registros y la edición.",
    )


def filter_dataframe_ui(df: pd.DataFrame, key: str) -> pd.DataFrame:
    if df.empty:
        return df
    c1, c2 = st.columns([2, 1])
    search = c1.text_input(
        "Buscar en la tabla",
        key=f"search_{key}",
        placeholder="Código, nombre, estado, responsable, descripción...",
    )
    if search:
        mask = df.astype(str).apply(
            lambda col: col.str.contains(search, case=False, na=False)
        ).any(axis=1)
        df = df[mask]

    if "estado_actual" in df.columns:
        states = sorted(df["estado_actual"].dropna().astype(str).unique().tolist())
        selected_states = c2.multiselect(
            "Filtrar por estado",
            states,
            key=f"states_{key}",
        )
        if selected_states:
            df = df[df["estado_actual"].astype(str).isin(selected_states)]
    elif "estado_revision" in df.columns:
        states = sorted(df["estado_revision"].dropna().astype(str).unique().tolist())
        selected_states = c2.multiselect(
            "Filtrar por estado",
            states,
            key=f"states_{key}",
        )
        if selected_states:
            df = df[df["estado_revision"].astype(str).isin(selected_states)]
    return df


def dataframe_selection(
    df: pd.DataFrame,
    key: str,
    id_column: str,
    columns: Optional[List[str]] = None,
) -> List[str]:
    if df.empty:
        st.info("No hay registros para mostrar.")
        return []

    visible_columns = [col for col in (columns or list(df.columns)) if col in df.columns]
    view = df[visible_columns].copy()

    selected_ids: List[str] = []
    try:
        event = st.dataframe(
            view,
            use_container_width=True,
            hide_index=True,
            key=f"grid_{key}",
            on_select="rerun",
            selection_mode="multi-row",
        )
        rows = event.selection.rows
        selected_ids = [
            str(df.iloc[position][id_column])
            for position in rows
            if position < len(df)
        ]
    except (TypeError, AttributeError):
        st.dataframe(view, use_container_width=True, hide_index=True)
        labels = {
            f"{row[id_column]} · {row.get('codigo_caso', row.get('tipo_actuacion', 'Registro'))}": str(row[id_column])
            for _, row in df.iterrows()
        }
        selected_labels = st.multiselect(
            "Seleccionar registros",
            list(labels.keys()),
            key=f"fallback_select_{key}",
        )
        selected_ids = [labels[label] for label in selected_labels]

    st.caption(
        f"Registros visibles: {len(df)} · Seleccionados: {len(selected_ids)}"
    )
    return selected_ids


def export_all_tables_excel() -> bytes:
    output = BytesIO()
    with pd.ExcelWriter(output, engine="openpyxl") as writer:
        with db_connection() as conn:
            for table_name in TABLE_EXPORT_ORDER:
                if not sql_table_exists(conn, table_name):
                    continue
                df = pd.read_sql_query(f'SELECT * FROM "{table_name}"', conn)
                sheet_name = table_name[:31]
                df.to_excel(writer, sheet_name=sheet_name, index=False)
    output.seek(0)
    return output.getvalue()


def export_dataframe_excel(df: pd.DataFrame, sheet_name: str) -> bytes:
    output = BytesIO()
    with pd.ExcelWriter(output, engine="openpyxl") as writer:
        df.to_excel(writer, sheet_name=sheet_name[:31], index=False)
    output.seek(0)
    return output.getvalue()


def format_date_es(value: Any, include_time: bool = False) -> str:
    if value in (None, ""):
        return ""
    parsed = pd.to_datetime(value, errors="coerce")
    if pd.isna(parsed):
        return str(value)
    if include_time:
        return parsed.strftime("%d/%b/%Y %I:%M %p")
    return parsed.strftime("%d/%b/%Y")




def format_time_12h(value: Any) -> str:
    if value in (None, ""):
        return ""
    text = str(value).strip()
    for fmt in ("%H:%M", "%H:%M:%S", "%I:%M %p", "%I:%M:%S %p"):
        try:
            return datetime.strptime(text, fmt).strftime("%I:%M %p")
        except ValueError:
            continue
    parsed = pd.to_datetime(text, errors="coerce")
    return text if pd.isna(parsed) else parsed.strftime("%I:%M %p")


def parse_time_to_24h(value: Any, default: str = "08:00") -> str:
    if value in (None, ""):
        return default
    text = str(value).strip().upper().replace("A. M.", "AM").replace("P. M.", "PM")
    for fmt in ("%H:%M", "%H:%M:%S", "%I:%M %p", "%I:%M:%S %p"):
        try:
            return datetime.strptime(text, fmt).strftime("%H:%M")
        except ValueError:
            continue
    parsed = pd.to_datetime(text, errors="coerce")
    return default if pd.isna(parsed) else parsed.strftime("%H:%M")


def time_select_12h(container, label: str, current_value: Optional[str], key: str) -> str:
    """Captura hora y minuto exactos en formato de 12 horas y almacena HH:MM."""
    current_24 = parse_time_to_24h(current_value, datetime.now().strftime("%H:%M"))
    parsed = datetime.strptime(current_24, "%H:%M")
    hour_12 = parsed.hour % 12 or 12
    period = "AM" if parsed.hour < 12 else "PM"

    container.markdown(f"**{label}**")
    c_hour, c_minute, c_period = container.columns([1, 1, 1])
    selected_hour = int(c_hour.number_input(
        "Hora", min_value=1, max_value=12, value=hour_12, step=1,
        key=f"{key}_hour",
    ))
    selected_minute = int(c_minute.number_input(
        "Minutos", min_value=0, max_value=59, value=parsed.minute, step=1,
        key=f"{key}_minute",
    ))
    selected_period = c_period.selectbox(
        "AM / PM", ["AM", "PM"], index=0 if period == "AM" else 1,
        key=f"{key}_period",
    )
    hour_24 = selected_hour % 12
    if selected_period == "PM":
        hour_24 += 12
    stored = f"{hour_24:02d}:{selected_minute:02d}"
    container.caption(f"Hora registrada: {format_time_12h(stored)}")
    return stored


def get_case_bundle(case_id: str):
    with db_connection() as conn:
        case = conn.execute(
            "SELECT * FROM casos WHERE id_caso = ?",
            (case_id,),
        ).fetchone()
        followups = conn.execute(
            """
            SELECT * FROM seguimientos
            WHERE id_caso = ?
            ORDER BY fecha_actuacion, fecha_registro_sistema
            """,
            (case_id,),
        ).fetchall()
    return case, followups


def add_word_field(table, row_index: int, values: List[str]) -> None:
    cells = table.rows[row_index].cells
    for idx, value in enumerate(values):
        if idx < len(cells):
            cells[idx].text = str(value or "")


def build_cases_word(case_ids: List[str]) -> bytes:
    document = Document()
    section = document.sections[0]
    section.top_margin = Cm(1.3)
    section.bottom_margin = Cm(1.3)
    section.left_margin = Cm(1.4)
    section.right_margin = Cm(1.4)

    for case_position, case_id in enumerate(case_ids):
        case, followups = get_case_bundle(case_id)
        if not case:
            continue

        heading = document.add_paragraph()
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = heading.add_run("Autoridad del Canal de Panamá\n")
        run.bold = True
        run.font.size = Pt(12)
        run = heading.add_run("Oficina de Proyectos Hídricos\n")
        run.bold = True
        run = heading.add_run("Formulario de casos del Programa Hídrico")
        run.bold = True

        header = document.add_table(rows=1, cols=2)
        header.style = "Table Grid"
        add_word_field(
            header,
            0,
            [
                f"ID M8: {case['id_control_m8']} · Caso: {case['codigo_caso']}",
                (
                    "Fecha de registro: "
                    f"{format_date_es(case['fecha_registro'])} "
                    f"{format_time_12h(case['hora_registro'] or case['hora_recepcion'])}"
                ),
            ],
        )

        document.add_heading("Información del Contacto", level=2)
        table = document.add_table(rows=2, cols=3)
        table.style = "Table Grid"
        add_word_field(table, 0, [
            f"Nombre: {case['nombre_solicitante'] or 'Sin nombre registrado'}",
            f"Sexo: {case['sexo'] or ''}",
            f"Cédula del contacto: {case['cedula'] or ''}",
        ])
        add_word_field(table, 1, [
            f"Teléfono: {case['telefono'] or ''}",
            f"Celular: {case['celular'] or ''}",
            f"Correo electrónico: {case['correo'] or ''}",
        ])

        document.add_heading("Ubicación del Contacto", level=2)
        table = document.add_table(rows=2, cols=4)
        table.style = "Table Grid"
        add_word_field(table, 0, [
            f"Provincia: {case['provincia_contacto'] or ''}",
            f"Distrito: {case['distrito_contacto'] or ''}",
            f"Corregimiento: {case['corregimiento_contacto'] or ''}",
            f"Comunidad: {case['lugar_poblado_contacto'] or ''}",
        ])
        merged = table.cell(1, 0).merge(table.cell(1, 3))
        merged.text = f"Dirección: {case['direccion_contacto'] or ''}"

        title = "Datos del Caso"
        document.add_heading(title, level=2)
        details = [
            ("Descripción", case["descripcion"]),
            ("Nombre del responsable del origen", case["responsable_origen"]),
            ("¿Ha sido presentada anteriormente?", "Sí" if case["presentado_anteriormente"] else "No"),
            ("Tema", case["tema"]),
            ("Medio por el que se recibió", case["medio_recepcion"]),
            ("Otro medio de recepción", case["otro_medio_recepcion"]),
            ("Recepción por", case["recibido_por"]),
            ("Tipo de caso", case["tipo_caso"] or case["tipo_queja"]),
            ("Respuesta inmediata", case["respuesta_inmediata"]),
            ("Propuesta de solución", case["propuesta_solucion"]),
        ]
        table = document.add_table(rows=len(details), cols=2)
        table.style = "Table Grid"
        for idx, (label, value) in enumerate(details):
            add_word_field(table, idx, [label, value])

        document.add_heading("Ubicación del Caso", level=2)
        table = document.add_table(rows=2, cols=4)
        table.style = "Table Grid"
        add_word_field(table, 0, [
            f"Provincia: {case['provincia_hecho'] or ''}",
            f"Distrito: {case['distrito_hecho'] or ''}",
            f"Corregimiento: {case['corregimiento_hecho'] or ''}",
            f"Comunidad: {case['lugar_poblado_hecho'] or ''}",
        ])
        merged = table.cell(1, 0).merge(table.cell(1, 3))
        merged.text = f"Dirección: {case['direccion_hecho'] or ''}"

        document.add_heading("Seguimiento del Caso", level=2)
        table = document.add_table(rows=1, cols=2)
        table.style = "Table Grid"
        add_word_field(table, 0, ["Fecha", "Descripción"])
        for followup in followups:
            cells = table.add_row().cells
            cells[0].text = (
                f"{format_date_es(followup['fecha_actuacion'])} "
                f"{format_time_12h(followup['hora_actuacion'])}"
            ).strip()
            cells[1].text = followup["descripcion"] or ""

        document.add_heading("Cierre del Caso", level=2)
        closing = [
            ("Respuesta brindada", case["respuesta_final"]),
            ("¿Acepta la respuesta?", case["acepta_respuesta"]),
            ("Nivel de satisfacción", case["nivel_satisfaccion"]),
            ("Comentarios finales", case["comentarios_finales"]),
            ("Firma", "____________________________"),
            (
                "Fecha",
                f"{format_date_es(case['fecha_cierre'])} {format_time_12h(case['hora_cierre'])}".strip(),
            ),
        ]
        table = document.add_table(rows=len(closing), cols=2)
        table.style = "Table Grid"
        for idx, (label, value) in enumerate(closing):
            add_word_field(table, idx, [label, value])

        if case_position < len(case_ids) - 1:
            document.add_page_break()

    output = BytesIO()
    document.save(output)
    output.seek(0)
    return output.getvalue()


def pdf_paragraph(value: Any, style):
    return Paragraph(escape(str(value or "")), style)


def build_cases_pdf(case_ids: List[str]) -> bytes:
    output = BytesIO()
    doc = SimpleDocTemplate(
        output,
        pagesize=A4,
        leftMargin=12 * mm,
        rightMargin=12 * mm,
        topMargin=10 * mm,
        bottomMargin=10 * mm,
    )
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        "case_title",
        parent=styles["Heading1"],
        fontName="Helvetica-Bold",
        fontSize=11,
        leading=14,
        alignment=TA_CENTER,
        textColor=colors.HexColor("#073B5A"),
    )
    section_style = ParagraphStyle(
        "case_section",
        parent=styles["Heading2"],
        fontName="Helvetica-Bold",
        fontSize=9.5,
        leading=12,
        textColor=colors.HexColor("#073B5A"),
        spaceBefore=5,
        spaceAfter=3,
    )
    cell_style = ParagraphStyle(
        "case_cell",
        parent=styles["Normal"],
        fontSize=7.6,
        leading=9.3,
        alignment=TA_LEFT,
    )
    label_style = ParagraphStyle(
        "case_label",
        parent=cell_style,
        fontName="Helvetica-Bold",
    )

    story = []
    for position, case_id in enumerate(case_ids):
        case, followups = get_case_bundle(case_id)
        if not case:
            continue

        story.append(pdf_paragraph(
            "Autoridad del Canal de Panamá<br/>"
            "Oficina de Proyectos Hídricos<br/>"
            "Formulario de casos del Programa Hídrico",
            title_style,
        ))
        story.append(Spacer(1, 5))

        header_data = [[
            pdf_paragraph(
                f"<b>ID M8:</b> {case['id_control_m8']} · <b>Caso:</b> {case['codigo_caso']}",
                cell_style,
            ),
            pdf_paragraph(
                f"<b>Fecha de registro:</b> {format_date_es(case['fecha_registro'])} "
                f"{format_time_12h(case['hora_registro'] or case['hora_recepcion'])}",
                cell_style,
            ),
        ]]
        header = Table(header_data, colWidths=[9 * cm, 9 * cm])
        header.setStyle(TableStyle([
            ("BOX", (0, 0), (-1, -1), .5, colors.HexColor("#D6DEE6")),
            ("VALIGN", (0, 0), (-1, -1), "TOP"),
            ("BACKGROUND", (0, 0), (-1, -1), colors.HexColor("#F4F7F9")),
            ("PADDING", (0, 0), (-1, -1), 5),
        ]))
        story.append(header)

        def section_table(title, pairs, widths=(5 * cm, 13 * cm)):
            story.append(pdf_paragraph(title, section_style))
            rows = []
            for label, value in pairs:
                rows.append([
                    pdf_paragraph(label, label_style),
                    pdf_paragraph(value, cell_style),
                ])
            table = Table(rows, colWidths=list(widths))
            table.setStyle(TableStyle([
                ("BOX", (0, 0), (-1, -1), .4, colors.HexColor("#D6DEE6")),
                ("INNERGRID", (0, 0), (-1, -1), .25, colors.HexColor("#E5EAF0")),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("PADDING", (0, 0), (-1, -1), 4),
            ]))
            story.append(table)

        section_table("Información del Contacto", [
            ("Nombre", case["nombre_solicitante"] or "Sin nombre registrado"),
            ("Sexo", case["sexo"]),
            ("Cédula", case["cedula"]),
            ("Teléfono / Celular", f"{case['telefono'] or ''} / {case['celular'] or ''}"),
            ("Correo electrónico", case["correo"]),
        ])
        section_table("Ubicación del Contacto", [
            ("Provincia", case["provincia_contacto"]),
            ("Distrito", case["distrito_contacto"]),
            ("Corregimiento", case["corregimiento_contacto"]),
            ("Comunidad", case["lugar_poblado_contacto"]),
            ("Dirección", case["direccion_contacto"]),
        ])
        section_table("Datos del Caso", [
            ("Descripción", case["descripcion"]),
            ("Responsable del origen", case["responsable_origen"]),
            ("Presentada anteriormente", "Sí" if case["presentado_anteriormente"] else "No"),
            ("Tema", case["tema"]),
            ("Medio de recepción", case["medio_recepcion"]),
            ("Otro medio", case["otro_medio_recepcion"]),
            ("Recepción por", case["recibido_por"]),
            ("Tipo de caso", case["tipo_caso"] or case["tipo_queja"]),
            ("Respuesta inmediata", case["respuesta_inmediata"]),
            ("Propuesta de solución", case["propuesta_solucion"]),
        ])
        section_table("Ubicación del Caso", [
            ("Provincia", case["provincia_hecho"]),
            ("Distrito", case["distrito_hecho"]),
            ("Corregimiento", case["corregimiento_hecho"]),
            ("Comunidad", case["lugar_poblado_hecho"]),
            ("Dirección", case["direccion_hecho"]),
        ])

        story.append(pdf_paragraph("Seguimiento del Caso", section_style))
        followup_rows = [[
            pdf_paragraph("Fecha", label_style),
            pdf_paragraph("Descripción", label_style),
        ]]
        for item in followups:
            followup_rows.append([
                pdf_paragraph(
                    f"{format_date_es(item['fecha_actuacion'])} {format_time_12h(item['hora_actuacion'])}".strip(),
                    cell_style,
                ),
                pdf_paragraph(item["descripcion"], cell_style),
            ])
        if len(followup_rows) == 1:
            followup_rows.append([
                pdf_paragraph("", cell_style),
                pdf_paragraph("Sin seguimientos registrados.", cell_style),
            ])
        table = Table(followup_rows, colWidths=[3 * cm, 15 * cm], repeatRows=1)
        table.setStyle(TableStyle([
            ("BOX", (0, 0), (-1, -1), .4, colors.HexColor("#D6DEE6")),
            ("INNERGRID", (0, 0), (-1, -1), .25, colors.HexColor("#E5EAF0")),
            ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#F4F7F9")),
            ("VALIGN", (0, 0), (-1, -1), "TOP"),
            ("PADDING", (0, 0), (-1, -1), 4),
        ]))
        story.append(table)

        section_table("Cierre del Caso", [
            ("Respuesta brindada", case["respuesta_final"]),
            ("Acepta la respuesta", case["acepta_respuesta"]),
            ("Satisfacción", case["nivel_satisfaccion"]),
            ("Comentarios finales", case["comentarios_finales"]),
            ("Firma", "____________________________"),
            (
                "Fecha",
                f"{format_date_es(case['fecha_cierre'])} {format_time_12h(case['hora_cierre'])}".strip(),
            ),
        ])

        if position < len(case_ids) - 1:
            story.append(PageBreak())

    doc.build(story)
    output.seek(0)
    return output.getvalue()


def render_export_buttons(
    df: pd.DataFrame,
    selected_case_ids: Optional[List[str]],
    key: str,
    sheet_name: str,
) -> None:
    st.markdown("#### Descargas")
    c1, c2, c3, c4 = st.columns(4)
    c1.download_button(
        "Excel de tabla visible",
        data=export_dataframe_excel(df, sheet_name),
        file_name=f"{sheet_name.lower().replace(' ', '_')}_visible.xlsx",
        mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        use_container_width=True,
        key=f"excel_visible_{key}",
    )
    c2.download_button(
        "Excel de todas las tablas",
        data=export_all_tables_excel(),
        file_name="modulo_casos_todas_las_tablas.xlsx",
        mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        use_container_width=True,
        key=f"excel_all_{key}",
    )

    case_ids = selected_case_ids or []
    c3.download_button(
        "Word de formularios",
        data=build_cases_word(case_ids) if case_ids else b"",
        file_name="formularios_casos.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        disabled=not bool(case_ids),
        use_container_width=True,
        key=f"word_cases_{key}",
        help="Selecciona uno o varios casos en la tabla.",
    )
    c4.download_button(
        "PDF de formularios",
        data=build_cases_pdf(case_ids) if case_ids else b"",
        file_name="formularios_casos.pdf",
        mime="application/pdf",
        disabled=not bool(case_ids),
        use_container_width=True,
        key=f"pdf_cases_{key}",
        help="Selecciona uno o varios casos en la tabla.",
    )


def cases_history_df() -> pd.DataFrame:
    return read_query_df(
        """
        SELECT
            id_caso,
            id_control_m8,
            codigo_caso,
            fecha_registro,
            hora_registro,
            clasificacion,
            COALESCE(NULLIF(tipo_caso, ''), tipo_queja) AS tipo_caso,
            nombre_solicitante,
            cedula,
            tema,
            medio_recepcion,
            provincia_hecho,
            distrito_hecho,
            lugar_poblado_hecho,
            estado_actual,
            responsable_principal,
            id_persona_m01,
            id_hogar_m01,
            pertenece_proyecto,
            estado_vinculacion_m01,
            fecha_cierre,
            fecha_ultima_actualizacion
        FROM casos
        ORDER BY fecha_registro DESC, codigo_caso DESC
        """
    )


def followups_history_df() -> pd.DataFrame:
    return read_query_df(
        """
        SELECT
            s.id_seguimiento,
            s.id_control_seguimiento,
            c.id_caso,
            c.codigo_caso,
            c.clasificacion,
            s.fecha_actuacion,
            s.hora_actuacion,
            s.tipo_actuacion,
            s.descripcion,
            s.resultado,
            s.responsable_ejecutor,
            s.estado_anterior,
            s.estado_posterior,
            s.proxima_accion,
            s.fecha_compromiso,
            s.estado_actividad,
            s.informo_solicitante,
            s.tipo_informacion_solicitante,
            s.medio_informacion_solicitante,
            s.resultado_contacto_solicitante,
            s.contenido_comunicado,
            s.usuario_registro,
            s.fecha_registro_sistema
        FROM seguimientos s
        JOIN casos c ON c.id_caso = s.id_caso
        ORDER BY s.fecha_actuacion DESC, s.fecha_registro_sistema DESC
        """
    )

def communications_history_df() -> pd.DataFrame:
    return read_query_df(
        """
        SELECT
            m.id_comunicacion,
            c.id_caso,
            c.codigo_caso,
            c.clasificacion,
            m.tipo_comunicacion,
            m.fecha,
            m.medio,
            m.destinatario,
            m.realizada_por,
            m.resultado_contacto,
            m.descripcion,
            m.es_acuse_recibo,
            m.es_comunicacion_avance,
            m.es_comunicacion_cierre,
            m.fecha_registro_sistema
        FROM comunicaciones m
        JOIN casos c ON c.id_caso = m.id_caso
        ORDER BY m.fecha DESC, m.fecha_registro_sistema DESC
        """
    )


def approvals_history_df() -> pd.DataFrame:
    return read_query_df(
        """
        SELECT
            id_caso,
            id_control_m8,
            codigo_caso,
            clasificacion,
            estado_actual,
            responsable_principal,
            recomendacion_cierre,
            fecha_recomendacion_cierre,
            recomendada_por,
            visto_bueno_supervisor,
            fecha_visto_bueno,
            observacion_visto_bueno,
            decision_supervisor,
            fecha_decision_supervisor,
            resultado_atencion,
            se_brindo_respuesta,
            respuesta_final,
            acepta_respuesta,
            nivel_satisfaccion,
            autoriza_divulgacion_datos,
            fecha_cierre,
            hora_cierre,
            cerrado_por,
            motivo_cierre
        FROM casos
        WHERE recomendacion_cierre IS NOT NULL
           OR visto_bueno_supervisor IS NOT NULL
           OR fecha_cierre IS NOT NULL
        ORDER BY COALESCE(fecha_cierre, fecha_visto_bueno, fecha_recomendacion_cierre) DESC
        """
    )

def reviews_history_df() -> pd.DataFrame:
    return read_query_df(
        """
        SELECT
            r.id_revision,
            c.id_caso,
            c.codigo_caso,
            c.clasificacion,
            r.tipo_revision,
            r.fecha_solicitud,
            r.solicitada_por,
            r.motivo,
            r.revisor,
            r.fecha_revision,
            r.decision,
            r.observaciones,
            r.estado_revision
        FROM revisiones r
        JOIN casos c ON c.id_caso = r.id_caso
        ORDER BY r.fecha_solicitud DESC
        """
    )


def selected_case_ids_from_related(
    df: pd.DataFrame,
    selected_record_ids: List[str],
    record_id_column: str,
) -> List[str]:
    if not selected_record_ids or df.empty:
        return []
    selected = df[df[record_id_column].astype(str).isin(selected_record_ids)]
    if "id_caso" not in selected.columns:
        return []
    return selected["id_caso"].dropna().astype(str).unique().tolist()


def case_form(defaults: Optional[sqlite3.Row] = None, form_key: str = "case_form"):
    values = dict(defaults) if defaults else {}

    with st.form(form_key):
        st.markdown("#### Identificación y registro del caso")
        c1, c2, c3 = st.columns(3)
        origen_values = _options_with_current(["ACP", "Contratista"], values.get("origen_gestion"), False)
        origen_gestion = c1.selectbox(
            "Origen de la gestión",
            origen_values,
            index=_option_index(origen_values, values.get("origen_gestion") or "ACP"),
        )
        fecha_registro = c2.date_input(
            "Fecha de registro",
            value=date.fromisoformat(values["fecha_registro"])
            if values.get("fecha_registro") else date.today(),
        )
        clasificacion_values = _options_with_current(
            CLASIFICACION_OPTIONS, values.get("clasificacion"), False
        )
        clasificacion = c3.selectbox(
            "Clasificación del caso",
            clasificacion_values,
            index=_option_index(
                clasificacion_values, values.get("clasificacion") or "Consulta"
            ),
        )
        hora_registro = time_select_12h(
            st,
            "Hora de recepción",
            values.get("hora_registro") or values.get("hora_recepcion"),
            key=f"{form_key}_hora_registro",
        )

        c1, c2 = st.columns(2)
        trabajador_registra = c1.text_input(
            "Nombre de quien registra",
            value=values.get("registrado_por") or st.session_state.current_user,
        )
        cargo_registra = c2.text_input(
            "Cargo de quien registra",
            value=values.get("cargo_registra") or "",
        )

        st.markdown("#### Información de la persona")
        c1, c2, c3 = st.columns(3)
        nombre = c1.text_input(
            "Nombre completo",
            value=values.get("nombre_solicitante") or "",
        )
        sexo_current = values.get("sexo") or ""
        sexo_aliases = {"Masculino": "Hombre", "Femenino": "Mujer"}
        sexo_current = sexo_aliases.get(sexo_current, sexo_current)
        sexo_values = _options_with_current(SEXO_OPTIONS, sexo_current)
        sexo = c2.selectbox(
            "Sexo",
            sexo_values,
            index=_option_index(sexo_values, sexo_current),
        )
        cedula = c3.text_input(
            "Cédula o pasaporte",
            value=values.get("cedula") or "",
        )

        c1, c2 = st.columns(2)
        contacto_options = _options_with_current(
            CONTACTO_PERSONA_OPTIONS, values.get("contacto_persona")
        )
        contacto_persona = c1.selectbox(
            "Contacto preferido de la persona",
            contacto_options,
            index=_option_index(contacto_options, values.get("contacto_persona") or ""),
        )
        telefono_celular_options = _options_with_current(
            TELEFONO_CELULAR_OPTIONS, values.get("telefono_o_celular")
        )
        telefono_o_celular = c2.selectbox(
            "Teléfono o celular",
            telefono_celular_options,
            index=_option_index(
                telefono_celular_options, values.get("telefono_o_celular") or ""
            ),
        )
        c1, c2, c3 = st.columns(3)
        telefono = c1.text_input("Teléfono", value=values.get("telefono") or "")
        celular = c2.text_input("Celular", value=values.get("celular") or "")
        correo = c3.text_input("Correo electrónico", value=values.get("correo") or "")

        st.markdown("#### Ubicación de la persona")
        c1, c2, c3 = st.columns(3)
        provincia_contacto_options = _location_options(
            PROVINCIAS,
            values.get("provincia_contacto_codigo"),
            values.get("provincia_contacto"),
        )
        provincia_contacto_option = c1.selectbox(
            "Provincia",
            provincia_contacto_options,
            index=_option_index(
                provincia_contacto_options,
                _location_option(
                    values.get("provincia_contacto_codigo"),
                    values.get("provincia_contacto"),
                ),
            ),
            key=f"{form_key}_prov_contacto",
        )
        distrito_contacto_options = _location_options(
            DISTRITOS,
            values.get("distrito_contacto_codigo"),
            values.get("distrito_contacto"),
        )
        distrito_contacto_option = c2.selectbox(
            "Distrito",
            distrito_contacto_options,
            index=_option_index(
                distrito_contacto_options,
                _location_option(
                    values.get("distrito_contacto_codigo"),
                    values.get("distrito_contacto"),
                ),
            ),
            key=f"{form_key}_dist_contacto",
        )
        corregimiento_contacto_options = _location_options(
            CORREGIMIENTOS,
            values.get("corregimiento_contacto_codigo"),
            values.get("corregimiento_contacto"),
        )
        corregimiento_contacto_option = c3.selectbox(
            "Corregimiento",
            corregimiento_contacto_options,
            index=_option_index(
                corregimiento_contacto_options,
                _location_option(
                    values.get("corregimiento_contacto_codigo"),
                    values.get("corregimiento_contacto"),
                ),
            ),
            key=f"{form_key}_corr_contacto",
        )
        c1, c2 = st.columns(2)
        comunidad_contacto = c1.text_input(
            "Comunidad",
            value=values.get("lugar_poblado_contacto") or "",
            help="El catálogo de comunidades proviene de un archivo externo del formulario y se mantiene como texto hasta recibirlo.",
        )
        comunidad_contacto_codigo = c2.text_input(
            "Código de comunidad",
            value=values.get("lugar_poblado_contacto_codigo") or "",
        )
        direccion_contacto = st.text_area(
            "Dirección completa",
            value=values.get("direccion_contacto") or "",
            height=110,
        )

        st.markdown("#### Recepción del caso")
        c1, c2, c3 = st.columns(3)
        medio_options = _options_with_current(
            MEDIOS_RECEPCION, values.get("medio_recepcion"), False
        )
        medio = c1.selectbox(
            "Medio por el que se recibió",
            medio_options,
            index=_option_index(
                medio_options, values.get("medio_recepcion") or MEDIOS_RECEPCION[0]
            ),
        )
        otro_medio = c2.text_input(
            "Otro medio de recepción, cuando aplique",
            value=values.get("otro_medio_recepcion") or "",
        )
        oficina_options = _options_with_current(
            OFICINAS_RECEPCION, values.get("recibido_por"), False
        )
        recibido_por = c3.selectbox(
            "Recepción del caso por",
            oficina_options,
            index=_option_index(
                oficina_options, values.get("recibido_por") or OFICINAS_RECEPCION[0]
            ),
        )

        st.markdown("#### Datos del caso")
        tema_options, tema_selected = _parse_multiselect(
            values.get("tema"), TEMAS_CASO
        )
        tema = st.multiselect(
            "Tema del caso",
            tema_options,
            default=tema_selected,
        )
        c1, c2 = st.columns(2)
        tipo_options = _options_with_current(
            TIPOS_CASO, (values.get("tipo_caso") or values.get("tipo_queja"))
        )
        tipo_caso = c1.selectbox(
            "Tipo de caso",
            tipo_options,
            index=_option_index(tipo_options, (values.get("tipo_caso") or values.get("tipo_queja")) or ""),
        )
        tipo_caso_otro = c2.text_area(
            "Especifique otro tipo, cuando aplique",
            value=(values.get("tipo_caso_otro") or values.get("tipo_queja_otro") or ""),
            height=90,
        )

        c1, c2 = st.columns(2)
        responsable_catalogo_options = _options_with_current(
            SI_NO_OPTIONS, values.get("responsable_origen_catalogo")
        )
        responsable_origen_catalogo = c1.selectbox(
            "¿Conoce al responsable del hecho o condición que origina el caso?",
            responsable_catalogo_options,
            index=_option_index(
                responsable_catalogo_options,
                values.get("responsable_origen_catalogo") or "",
            ),
        )
        nombre_responsable_origen = c2.text_input(
            "Nombre del responsable, cuando se conozca",
            value=values.get("nombre_responsable_origen")
            or values.get("responsable_origen")
            or "",
        )

        descripcion = st.text_area(
            "Descripción del caso",
            value=values.get("descripcion") or "",
            height=190,
        )
        presented = st.checkbox(
            "¿El caso ha sido presentado anteriormente?",
            value=bool(values.get("presentado_anteriormente", 0)),
        )
        previous_reference = st.text_area(
            "Si la respuesta es afirmativa, explique o indique la referencia anterior",
            value=values.get("referencia_caso_anterior") or "",
            height=120,
            help="Este campo se exige al guardar cuando la opción anterior está marcada.",
        )
        respuesta_inmediata = st.text_area(
            "Respuesta inmediata al caso",
            value=values.get("respuesta_inmediata") or "",
            height=140,
        )
        propuesta = st.text_area(
            "Propuesta de solución del caso",
            value=values.get("propuesta_solucion") or "",
            height=150,
        )

        st.markdown("#### Ubicación del caso")
        c1, c2, c3 = st.columns(3)
        provincia_hecho_options = _location_options(
            PROVINCIAS, values.get("provincia_hecho_codigo"), values.get("provincia_hecho")
        )
        provincia_hecho_option = c1.selectbox(
            "Provincia del caso",
            provincia_hecho_options,
            index=_option_index(
                provincia_hecho_options,
                _location_option(values.get("provincia_hecho_codigo"), values.get("provincia_hecho")),
            ),
            key=f"{form_key}_prov_hecho",
        )
        distrito_hecho_options = _location_options(
            DISTRITOS, values.get("distrito_hecho_codigo"), values.get("distrito_hecho")
        )
        distrito_hecho_option = c2.selectbox(
            "Distrito del caso",
            distrito_hecho_options,
            index=_option_index(
                distrito_hecho_options,
                _location_option(values.get("distrito_hecho_codigo"), values.get("distrito_hecho")),
            ),
            key=f"{form_key}_dist_hecho",
        )
        corregimiento_hecho_options = _location_options(
            CORREGIMIENTOS,
            values.get("corregimiento_hecho_codigo"),
            values.get("corregimiento_hecho"),
        )
        corregimiento_hecho_option = c3.selectbox(
            "Corregimiento del caso",
            corregimiento_hecho_options,
            index=_option_index(
                corregimiento_hecho_options,
                _location_option(
                    values.get("corregimiento_hecho_codigo"),
                    values.get("corregimiento_hecho"),
                ),
            ),
            key=f"{form_key}_corr_hecho",
        )
        c1, c2 = st.columns(2)
        comunidad_hecho = c1.text_input(
            "Comunidad del caso",
            value=values.get("lugar_poblado_hecho") or "",
        )
        comunidad_hecho_codigo = c2.text_input(
            "Código de comunidad del caso",
            value=values.get("lugar_poblado_hecho_codigo") or "",
        )
        direccion_hecho = st.text_area(
            "Dirección o ubicación exacta del caso",
            value=values.get("direccion_hecho") or "",
            height=130,
        )

        referencia_options, referencia_selected = _parse_multiselect(
            values.get("referencia_afectacion"), REFERENCIAS_AFECTACION
        )
        referencia_afectacion = st.multiselect(
            "El lugar del caso está cerca de",
            referencia_options,
            default=referencia_selected,
        )
        otra_referencia_afectacion = st.text_area(
            "Explique otra referencia, cuando aplique",
            value=values.get("otra_referencia_afectacion") or "",
            height=100,
        )
        llegada_lugar_afectacion = st.text_area(
            "¿Cómo se llega al lugar del caso?",
            value=values.get("llegada_lugar_afectacion") or "",
            height=120,
        )
        c1, c2, c3 = st.columns(3)
        coordenadas_options = _options_with_current(
            SI_NO_OPTIONS, values.get("tiene_coordenadas_hecho")
        )
        tiene_coordenadas_hecho = c1.selectbox(
            "¿Cuenta con coordenadas del sitio?",
            coordenadas_options,
            index=_option_index(
                coordenadas_options, values.get("tiene_coordenadas_hecho") or ""
            ),
        )
        coordenada_norte_utm = c2.number_input(
            "Coordenada Norte UTM",
            value=float(values.get("coordenada_norte_utm") or 0.0),
            format="%.6f",
        )
        coordenada_este_utm = c3.number_input(
            "Coordenada Este UTM",
            value=float(values.get("coordenada_este_utm") or 0.0),
            format="%.6f",
        )

        st.markdown("#### Asignación interna")
        c1, c2, c3 = st.columns(3)
        equipo_supervisor_options = _options_with_current(
            EQUIPOS_PH, values.get("equipo_supervisor")
        )
        equipo_supervisor = c1.selectbox(
            "Equipo supervisor",
            equipo_supervisor_options,
            index=_option_index(
                equipo_supervisor_options, values.get("equipo_supervisor") or ""
            ),
        )
        supervisor = c2.text_input(
            "Supervisor",
            value=values.get("supervisor") or "",
            help="El catálogo nominal de supervisores proviene de un archivo externo y se mantiene como texto hasta recibirlo.",
        )
        fecha_asignacion = c3.date_input(
            "Fecha de asignación",
            value=date.fromisoformat(values["fecha_asignacion"])
            if values.get("fecha_asignacion") else None,
        )
        cantidad_current = int(values.get("cantidad_responsables_asignacion") or 1)
        cantidad_responsables = st.selectbox(
            "Cantidad de responsables de la asignación",
            CANTIDAD_RESPONSABLES_OPTIONS,
            index=CANTIDAD_RESPONSABLES_OPTIONS.index(cantidad_current)
            if cantidad_current in CANTIDAD_RESPONSABLES_OPTIONS else 0,
        )
        c1, c2 = st.columns(2)
        responsable_1 = c1.text_input(
            "Responsable 1",
            value=values.get("responsable_principal") or "",
        )
        equipo_1_options = _options_with_current(
            EQUIPOS_PH, values.get("equipo_responsable_1")
        )
        equipo_1 = c2.selectbox(
            "Equipo responsable 1",
            equipo_1_options,
            index=_option_index(equipo_1_options, values.get("equipo_responsable_1") or ""),
        )
        c1, c2 = st.columns(2)
        responsable_2 = c1.text_input(
            "Responsable 2",
            value=values.get("responsable_apoyo_1") or "",
        )
        equipo_2_options = _options_with_current(
            EQUIPOS_PH, values.get("equipo_responsable_2")
        )
        equipo_2 = c2.selectbox(
            "Equipo responsable 2",
            equipo_2_options,
            index=_option_index(equipo_2_options, values.get("equipo_responsable_2") or ""),
        )
        c1, c2 = st.columns(2)
        responsable_3 = c1.text_input(
            "Responsable 3",
            value=values.get("responsable_apoyo_2") or "",
        )
        equipo_3_options = _options_with_current(
            EQUIPOS_PH, values.get("equipo_responsable_3")
        )
        equipo_3 = c2.selectbox(
            "Equipo responsable 3",
            equipo_3_options,
            index=_option_index(equipo_3_options, values.get("equipo_responsable_3") or ""),
        )

        files = st.file_uploader(
            "Adjuntos o evidencias del caso",
            accept_multiple_files=True,
            key=f"{form_key}_files",
        )
        submitted = st.form_submit_button(
            "Guardar caso",
            type="primary",
            use_container_width=True,
        )

    provincia_contacto_codigo, provincia_contacto = _decode_location_option(
        provincia_contacto_option
    )
    distrito_contacto_codigo, distrito_contacto = _decode_location_option(
        distrito_contacto_option
    )
    corregimiento_contacto_codigo, corregimiento_contacto = _decode_location_option(
        corregimiento_contacto_option
    )
    provincia_hecho_codigo, provincia_hecho = _decode_location_option(
        provincia_hecho_option
    )
    distrito_hecho_codigo, distrito_hecho = _decode_location_option(
        distrito_hecho_option
    )
    corregimiento_hecho_codigo, corregimiento_hecho = _decode_location_option(
        corregimiento_hecho_option
    )

    payload = {
        "fecha_registro": fecha_registro.isoformat(),
        "hora_registro": hora_registro,
        "fecha_recepcion": fecha_registro.isoformat(),
        "hora_recepcion": hora_registro,
        "clasificacion": clasificacion,
        "origen_gestion": origen_gestion,
        "cargo_registra": normalize_text(cargo_registra),
        "respecto_a": "",
        "respecto_otro": "",
        "tipo_identificacion_solicitante": "Identificado",
        "confidencial": int(values.get("confidencial", 0)),
        "nombre_solicitante": normalize_text(nombre),
        "sexo": sexo,
        "cedula": normalize_text(cedula),
        "contacto_persona": contacto_persona,
        "telefono_o_celular": telefono_o_celular,
        "telefono": normalize_text(telefono),
        "celular": normalize_text(celular),
        "correo": normalize_text(correo),
        "provincia_contacto_codigo": provincia_contacto_codigo,
        "distrito_contacto_codigo": distrito_contacto_codigo,
        "corregimiento_contacto_codigo": corregimiento_contacto_codigo,
        "lugar_poblado_contacto_codigo": normalize_text(comunidad_contacto_codigo),
        "provincia_contacto": provincia_contacto,
        "distrito_contacto": distrito_contacto,
        "corregimiento_contacto": corregimiento_contacto,
        "lugar_poblado_contacto": normalize_text(comunidad_contacto),
        "direccion_contacto": normalize_text(direccion_contacto),
        "medio_recepcion": medio,
        "otro_medio_recepcion": normalize_text(otro_medio) if medio == "Otro" else "",
        "recibido_por": recibido_por,
        "tema": "; ".join(tema),
        "tipo_caso": tipo_caso,
        "tipo_caso_otro": normalize_text(tipo_caso_otro) if tipo_caso == "Otro" else "",
        # Se mantienen los campos históricos para no romper bases ni exportaciones anteriores.
        "tipo_queja": tipo_caso,
        "tipo_queja_otro": normalize_text(tipo_caso_otro) if tipo_caso == "Otro" else "",
        "responsable_origen_catalogo": responsable_origen_catalogo,
        "nombre_responsable_origen": normalize_text(nombre_responsable_origen)
        if responsable_origen_catalogo == "Sí" else "",
        "responsable_origen": normalize_text(nombre_responsable_origen)
        if responsable_origen_catalogo == "Sí" else responsable_origen_catalogo,
        "descripcion": normalize_text(descripcion),
        "presentado_anteriormente": int(presented),
        "referencia_caso_anterior": normalize_text(previous_reference) if presented else "",
        "respuesta_inmediata": normalize_text(respuesta_inmediata),
        "propuesta_solucion": normalize_text(propuesta),
        "provincia_hecho_codigo": provincia_hecho_codigo,
        "distrito_hecho_codigo": distrito_hecho_codigo,
        "corregimiento_hecho_codigo": corregimiento_hecho_codigo,
        "lugar_poblado_hecho_codigo": normalize_text(comunidad_hecho_codigo),
        "provincia_hecho": provincia_hecho,
        "distrito_hecho": distrito_hecho,
        "corregimiento_hecho": corregimiento_hecho,
        "lugar_poblado_hecho": normalize_text(comunidad_hecho),
        "direccion_hecho": normalize_text(direccion_hecho),
        "referencia_afectacion": "; ".join(referencia_afectacion),
        "otra_referencia_afectacion": normalize_text(otra_referencia_afectacion)
        if "Otra referencia" in referencia_afectacion else "",
        "llegada_lugar_afectacion": normalize_text(llegada_lugar_afectacion),
        "tiene_coordenadas_hecho": tiene_coordenadas_hecho,
        "coordenada_norte_utm": coordenada_norte_utm
        if tiene_coordenadas_hecho == "Sí" else None,
        "coordenada_este_utm": coordenada_este_utm
        if tiene_coordenadas_hecho == "Sí" else None,
        "equipo_supervisor": equipo_supervisor,
        "supervisor": normalize_text(supervisor),
        "fecha_asignacion": fecha_asignacion.isoformat() if fecha_asignacion else None,
        "cantidad_responsables_asignacion": cantidad_responsables,
        "responsable_principal": normalize_text(responsable_1),
        "equipo_responsable_1": equipo_1,
        "responsable_apoyo_1": normalize_text(responsable_2)
        if cantidad_responsables >= 2 else "",
        "equipo_responsable_2": equipo_2 if cantidad_responsables >= 2 else "",
        "responsable_apoyo_2": normalize_text(responsable_3)
        if cantidad_responsables >= 3 else "",
        "equipo_responsable_3": equipo_3 if cantidad_responsables >= 3 else "",
        "asignado_por": normalize_text(trabajador_registra),
    }
    return submitted, payload, files

def save_new_case(payload: Dict[str, Any], files) -> str:
    required_text(payload["descripcion"], "Descripción")
    required_text(payload["recibido_por"], "Recepción por")
    required_text(payload["nombre_solicitante"], "Nombre")
    if payload.get("presentado_anteriormente"):
        required_text(payload.get("referencia_caso_anterior", ""), "Referencia del caso presentado anteriormente")
    if payload.get("medio_recepcion") == "Otro":
        required_text(payload.get("otro_medio_recepcion", ""), "Otro medio de recepción")
    required_text(payload.get("tipo_caso", ""), "Tipo de caso")
    if payload.get("tipo_caso") == "Otro":
        required_text(payload.get("tipo_caso_otro", ""), "Especificación de otro tipo de caso")
    validate_required_case_location(payload)

    with db_connection() as conn:
        case_id = generate_id("CASO")
        year = int(payload["fecha_registro"][:4])
        case_code = generate_case_code(conn, year)
        now = now_iso()
        relation = datos_vinculacion_m01(conn, payload.get("cedula", ""))
        data = {
            "id_caso": case_id,
            "codigo_caso": case_code,
            **payload,
            **relation,
            "registrado_por": st.session_state.current_user,
            "estado_actual": "Pendiente de asignación",
            "situacion_atencion": "",
            "dentro_alcance": 1,
            "motivo_fuera_alcance": "",
            "fecha_creacion": now,
            "fecha_ultima_actualizacion": now,
            "actualizado_por": st.session_state.current_user,
            "fuente_registro": "Manual",
            "archivo_origen": "",
        }
        fields = ", ".join(data.keys())
        placeholders = ", ".join(["?"] * len(data))
        conn.execute(
            f"INSERT INTO casos ({fields}) VALUES ({placeholders})",
            tuple(data.values()),
        )
        _assign_control_code(
            conn, "casos", "id_caso", case_id, "id_control_m8", "M8-C-"
        )
        register_state_change(
            conn,
            case_id,
            None,
            "Pendiente de asignación",
            st.session_state.current_user,
            "Registro inicial del caso.",
        )
        audit_change(
            conn,
            case_id,
            "CREACIÓN DE CASO",
            st.session_state.current_user,
            new_data=data,
            details="Creación desde el formulario oficial.",
        )

    save_uploaded_files(
        case_id,
        files,
        "Evidencia inicial",
        st.session_state.current_user,
    )
    return case_code


def update_existing_case(case_id: str, payload: Dict[str, Any], files) -> None:
    required_text(payload["descripcion"], "Descripción")
    required_text(payload["recibido_por"], "Recepción por")
    if payload.get("presentado_anteriormente"):
        required_text(payload.get("referencia_caso_anterior", ""), "Referencia del caso presentado anteriormente")
    if payload.get("medio_recepcion") == "Otro":
        required_text(payload.get("otro_medio_recepcion", ""), "Otro medio de recepción")
    required_text(payload.get("tipo_caso", ""), "Tipo de caso")
    if payload.get("tipo_caso") == "Otro":
        required_text(payload.get("tipo_caso_otro", ""), "Especificación de otro tipo de caso")
    validate_required_case_location(payload)
    with db_connection() as conn:
        payload = {
            **payload,
            **datos_vinculacion_m01(conn, payload.get("cedula", "")),
        }
        update_case_fields(
            conn,
            case_id,
            payload,
            st.session_state.current_user,
            "EDICIÓN DE FICHA DE REGISTRO",
            "Actualización de la ficha oficial del caso.",
        )
    save_uploaded_files(
        case_id,
        files,
        "Evidencia adicional del registro",
        st.session_state.current_user,
    )


def render_cases_history():
    df = filter_dataframe_ui(cases_history_df(), "cases")
    columns = [
        "id_control_m8", "codigo_caso", "fecha_registro", "hora_registro", "clasificacion",
        "tipo_caso", "nombre_solicitante", "cedula", "tema", "estado_actual",
        "id_persona_m01", "id_hogar_m01", "pertenece_proyecto",
        "estado_vinculacion_m01", "responsable_principal",
        "provincia_hecho", "distrito_hecho", "lugar_poblado_hecho", "fecha_cierre",
    ]
    selected = dataframe_selection(df, "cases", "id_caso", columns)
    render_export_buttons(df, selected, "cases", "Casos")


def render_followups_history():
    df = filter_dataframe_ui(followups_history_df(), "followups")
    columns = [
        "id_control_seguimiento", "codigo_caso", "fecha_actuacion", "hora_actuacion",
        "tipo_actuacion", "descripcion", "resultado", "responsable_ejecutor",
        "proxima_accion", "fecha_compromiso", "estado_actividad",
        "informo_solicitante", "tipo_informacion_solicitante",
        "medio_informacion_solicitante", "resultado_contacto_solicitante",
        "contenido_comunicado",
    ]
    selected_records = dataframe_selection(
        df, "followups", "id_seguimiento", columns
    )
    case_ids = selected_case_ids_from_related(
        df, selected_records, "id_seguimiento"
    )
    render_export_buttons(df, case_ids, "followups", "Seguimientos")

def render_communications_history():
    df = filter_dataframe_ui(communications_history_df(), "communications")
    columns = [
        "codigo_caso", "tipo_comunicacion", "fecha", "medio",
        "destinatario", "realizada_por", "resultado_contacto", "descripcion",
    ]
    selected_records = dataframe_selection(
        df, "communications", "id_comunicacion", columns
    )
    case_ids = selected_case_ids_from_related(
        df, selected_records, "id_comunicacion"
    )
    render_export_buttons(df, case_ids, "communications", "Comunicaciones")


def render_approvals_history():
    df = filter_dataframe_ui(approvals_history_df(), "approvals")
    columns = [
        "id_control_m8", "codigo_caso", "clasificacion", "estado_actual",
        "responsable_principal", "recomendacion_cierre", "fecha_recomendacion_cierre",
        "recomendada_por", "visto_bueno_supervisor", "fecha_visto_bueno",
        "resultado_atencion", "se_brindo_respuesta", "acepta_respuesta",
        "nivel_satisfaccion", "autoriza_divulgacion_datos", "fecha_cierre",
        "hora_cierre", "cerrado_por",
    ]
    selected = dataframe_selection(df, "approvals", "id_caso", columns)
    render_export_buttons(df, selected, "approvals", "Cierres")

def render_reviews_history():
    df = filter_dataframe_ui(reviews_history_df(), "reviews")
    columns = [
        "codigo_caso", "clasificacion", "tipo_revision",
        "fecha_solicitud", "solicitada_por", "motivo",
        "revisor", "fecha_revision", "decision", "estado_revision",
    ]
    selected_records = dataframe_selection(df, "reviews", "id_revision", columns)
    case_ids = selected_case_ids_from_related(df, selected_records, "id_revision")
    render_export_buttons(df, case_ids, "reviews", "Revisiones")


def render_followup_form(edit_id: Optional[str] = None):
    existing = None
    if edit_id:
        with db_connection() as conn:
            existing = conn.execute(
                "SELECT * FROM seguimientos WHERE id_seguimiento = ?",
                (edit_id,),
            ).fetchone()

    case_id = existing["id_caso"] if existing else select_case_id(
        "Caso asociado", f"followup_case_{edit_id or 'new'}"
    )
    if not case_id:
        return

    with db_connection() as conn:
        case = get_case(conn, case_id)
    values = dict(existing) if existing else {}

    with st.form(f"followup_form_{edit_id or 'new'}"):
        c1, c2 = st.columns(2)
        action_date = c1.date_input(
            "Fecha de la actuación",
            value=date.fromisoformat(values["fecha_actuacion"])
            if values.get("fecha_actuacion") else date.today(),
        )
        action_time = time_select_12h(
            c2,
            "Hora de la actuación",
            values.get("hora_actuacion"),
            key=f"followup_time_{edit_id or 'new'}",
        )
        c1, c2 = st.columns(2)
        action_type_values = _options_with_current(
            FOLLOWUP_TYPES, values.get("tipo_actuacion"), False
        )
        action_type = c1.selectbox(
            "Tipo de actuación",
            action_type_values,
            index=_option_index(
                action_type_values, values.get("tipo_actuacion") or FOLLOWUP_TYPES[0]
            ),
        )
        executor = c2.text_input(
            "Responsable ejecutor",
            value=values.get("responsable_ejecutor") or case["responsable_principal"] or "",
        )
        description = st.text_area(
            "Descripción del seguimiento del caso",
            value=values.get("descripcion") or "",
            height=160,
        )
        result = st.text_area(
            "Resultado obtenido",
            value=values.get("resultado") or "",
            height=120,
        )
        c1, c2 = st.columns(2)
        state_current = values.get("estado_posterior") or case["estado_actual"]
        posterior_state = c1.selectbox(
            "Estado posterior",
            CASE_STATES,
            index=CASE_STATES.index(state_current) if state_current in CASE_STATES else 0,
        )
        activity_states = ["Pendiente", "En proceso", "Completada", "Cancelada"]
        activity_current = values.get("estado_actividad") or "Pendiente"
        activity_status = c2.selectbox(
            "Estado de la actividad",
            activity_states,
            index=activity_states.index(activity_current)
            if activity_current in activity_states else 0,
        )
        next_action = st.text_area(
            "Próxima acción",
            value=values.get("proxima_accion") or "",
            height=110,
        )
        commitment_date = st.date_input(
            "Fecha compromiso",
            value=date.fromisoformat(values["fecha_compromiso"])
            if values.get("fecha_compromiso") else None,
        )

        st.markdown("#### Información proporcionada a la persona")
        informed = st.checkbox(
            "Se informó a la persona sobre la recepción, avance o resultado",
            value=bool(values.get("informo_solicitante", values.get("visible_solicitante", 0))),
        )
        c1, c2, c3 = st.columns(3)
        information_type_options = _options_with_current(
            COMMUNICATION_TYPES, values.get("tipo_informacion_solicitante"), False
        )
        information_type = c1.selectbox(
            "Tipo de información",
            information_type_options,
            index=_option_index(
                information_type_options,
                values.get("tipo_informacion_solicitante") or COMMUNICATION_TYPES[0],
            ),
        )
        medium_options = _options_with_current(
            MEDIOS_RECEPCION, values.get("medio_informacion_solicitante"), False
        )
        information_medium = c2.selectbox(
            "Medio utilizado",
            medium_options,
            index=_option_index(
                medium_options,
                values.get("medio_informacion_solicitante") or MEDIOS_RECEPCION[0],
            ),
        )
        contact_results = [
            "Recibido", "No localizado", "Sin respuesta",
            "Datos de contacto no disponibles", "Rechazó la comunicación", "Otro",
        ]
        result_options = _options_with_current(
            contact_results, values.get("resultado_contacto_solicitante"), False
        )
        contact_result = c3.selectbox(
            "Resultado del contacto",
            result_options,
            index=_option_index(
                result_options,
                values.get("resultado_contacto_solicitante") or contact_results[0],
            ),
        )
        communicated_content = st.text_area(
            "Contenido comunicado",
            value=values.get("contenido_comunicado") or "",
            height=130,
        )

        files = st.file_uploader(
            "Adjuntos del seguimiento",
            accept_multiple_files=True,
            key=f"followup_files_{edit_id or 'new'}",
        )
        save = st.form_submit_button(
            "Guardar seguimiento",
            type="primary",
            use_container_width=True,
        )

    if save:
        try:
            required_text(executor, "Responsable ejecutor")
            required_text(description, "Descripción del seguimiento")
            if informed:
                required_text(communicated_content, "Contenido comunicado")
            if action_date < date.fromisoformat(case["fecha_registro"]):
                raise ValueError("La actuación no puede ser anterior al registro del caso.")
            if commitment_date and commitment_date < action_date:
                raise ValueError("La fecha compromiso no puede ser anterior a la actuación.")

            values_to_save = {
                "id_caso": case_id,
                "fecha_actuacion": action_date.isoformat(),
                "hora_actuacion": action_time,
                "tipo_actuacion": action_type,
                "descripcion": normalize_text(description),
                "resultado": normalize_text(result),
                "responsable_ejecutor": normalize_text(executor),
                "usuario_registro": st.session_state.current_user,
                "estado_anterior": values.get("estado_anterior") or case["estado_actual"],
                "estado_posterior": posterior_state,
                "proxima_accion": normalize_text(next_action),
                "fecha_compromiso": commitment_date.isoformat() if commitment_date else None,
                "estado_actividad": activity_status,
                "visible_solicitante": int(informed),
                "informo_solicitante": int(informed),
                "tipo_informacion_solicitante": information_type if informed else "",
                "medio_informacion_solicitante": information_medium if informed else "",
                "resultado_contacto_solicitante": contact_result if informed else "",
                "contenido_comunicado": normalize_text(communicated_content) if informed else "",
                "fecha_registro_sistema": values.get("fecha_registro_sistema") or now_iso(),
            }

            with db_connection() as conn:
                if edit_id:
                    existing_row = dict(existing)
                    _update(
                        conn,
                        "seguimientos",
                        "id_seguimiento",
                        edit_id,
                        values_to_save,
                    )
                    followup_id = edit_id
                    audit_change(
                        conn,
                        case_id,
                        "EDICIÓN DE SEGUIMIENTO",
                        st.session_state.current_user,
                        previous_data=existing_row,
                        new_data=values_to_save,
                        details=description,
                    )
                else:
                    followup_id = generate_id("SEG")
                    columns = ["id_seguimiento"] + list(values_to_save.keys())
                    insert_values = [followup_id] + list(values_to_save.values())
                    conn.execute(
                        f"INSERT INTO seguimientos ({', '.join(columns)}) "
                        f"VALUES ({', '.join(['?'] * len(columns))})",
                        insert_values,
                    )
                    _assign_control_code(
                        conn, "seguimientos", "id_seguimiento", followup_id,
                        "id_control_seguimiento", "M8-S-"
                    )
                    audit_change(
                        conn,
                        case_id,
                        "NUEVO SEGUIMIENTO",
                        st.session_state.current_user,
                        new_data=values_to_save,
                        details=description,
                    )

                if posterior_state != case["estado_actual"]:
                    update_case_fields(
                        conn,
                        case_id,
                        {"estado_actual": posterior_state},
                        st.session_state.current_user,
                        "CAMBIO DE ESTADO POR SEGUIMIENTO",
                        description,
                    )
                    register_state_change(
                        conn,
                        case_id,
                        case["estado_actual"],
                        posterior_state,
                        st.session_state.current_user,
                        description,
                        followup_id,
                    )

            if files:
                case_dir = UPLOAD_DIR / case_id
                case_dir.mkdir(exist_ok=True)
                with db_connection() as conn:
                    for uploaded in files:
                        safe_name = Path(uploaded.name).name
                        target = case_dir / f"{uuid.uuid4().hex[:8]}_{safe_name}"
                        target.write_bytes(uploaded.getbuffer())
                        conn.execute(
                            """
                            INSERT INTO documentos (
                                id_documento, id_caso, id_seguimiento,
                                tipo_documento, nombre_archivo,
                                ruta_archivo, fecha_carga, cargado_por
                            ) VALUES (?, ?, ?, ?, ?, ?, ?, ?)
                            """,
                            (
                                generate_id("DOC"), case_id, followup_id,
                                "Evidencia de seguimiento", safe_name,
                                str(target), now_iso(), st.session_state.current_user,
                            ),
                        )
            st.success("Seguimiento guardado correctamente.")
            st.rerun()
        except Exception as exc:
            st.error(str(exc))

def render_communication_form(edit_id: Optional[str] = None):
    existing = None
    if edit_id:
        with db_connection() as conn:
            existing = conn.execute(
                "SELECT * FROM comunicaciones WHERE id_comunicacion = ?",
                (edit_id,),
            ).fetchone()

    case_id = existing["id_caso"] if existing else select_case_id(
        "Caso asociado", f"communication_case_{edit_id or 'new'}"
    )
    if not case_id:
        return
    with db_connection() as conn:
        case = get_case(conn, case_id)
    values = dict(existing) if existing else {}

    with st.form(f"communication_form_{edit_id or 'new'}"):
        c1, c2 = st.columns(2)
        type_current = values.get("tipo_comunicacion") or COMMUNICATION_TYPES[0]
        comm_type = c1.selectbox(
            "Tipo de comunicación",
            COMMUNICATION_TYPES,
            index=COMMUNICATION_TYPES.index(type_current) if type_current in COMMUNICATION_TYPES else 0,
        )
        comm_date = c2.date_input(
            "Fecha",
            value=date.fromisoformat(values["fecha"]) if values.get("fecha") else date.today(),
        )
        c1, c2 = st.columns(2)
        medium_current = values.get("medio") or RECEPTION_CHANNELS[0]
        medium = c1.selectbox(
            "Medio",
            RECEPTION_CHANNELS,
            index=RECEPTION_CHANNELS.index(medium_current) if medium_current in RECEPTION_CHANNELS else 0,
        )
        recipient = c2.text_input(
            "Destinatario",
            value=values.get("destinatario") or case["nombre_solicitante"] or "Solicitante",
        )
        c1, c2 = st.columns(2)
        performed_by = c1.text_input(
            "Realizada por",
            value=values.get("realizada_por") or st.session_state.current_user,
        )
        result_values = [
            "Recibido", "No localizado", "Sin respuesta",
            "Datos de contacto no disponibles",
            "Rechazó la comunicación", "Otro",
        ]
        result_current = values.get("resultado_contacto") or result_values[0]
        result_contact = c2.selectbox(
            "Resultado del contacto",
            result_values,
            index=result_values.index(result_current) if result_current in result_values else 0,
        )
        description = st.text_area(
            "Descripción o contenido comunicado",
            value=values.get("descripcion") or "",
            height=130,
        )
        files = st.file_uploader(
            "Evidencia de la comunicación",
            accept_multiple_files=True,
            key=f"communication_files_{edit_id or 'new'}",
        )
        save = st.form_submit_button(
            "Guardar comunicación",
            type="primary",
            use_container_width=True,
        )

    if save:
        try:
            required_text(performed_by, "Realizada por")
            required_text(description, "Descripción")
            values_to_save = {
                "id_caso": case_id,
                "tipo_comunicacion": comm_type,
                "fecha": comm_date.isoformat(),
                "hora": None,
                "medio": medium,
                "destinatario": recipient,
                "realizada_por": performed_by,
                "resultado_contacto": result_contact,
                "descripcion": description,
                "es_acuse_recibo": int(comm_type == "Acuse de recibo"),
                "es_comunicacion_avance": int(comm_type == "Comunicación de avance"),
                "es_comunicacion_cierre": int(
                    comm_type in ("Comunicación de resultado", "Comunicación de cierre")
                ),
                "fecha_registro_sistema": values.get("fecha_registro_sistema") or now_iso(),
            }
            with db_connection() as conn:
                if edit_id:
                    _update(
                        conn,
                        "comunicaciones",
                        "id_comunicacion",
                        edit_id,
                        values_to_save,
                    )
                    audit_change(
                        conn,
                        case_id,
                        "EDICIÓN DE COMUNICACIÓN",
                        st.session_state.current_user,
                        previous_data=dict(existing),
                        new_data=values_to_save,
                        details=description,
                    )
                    communication_id = edit_id
                else:
                    communication_id = generate_id("COM")
                    columns = ["id_comunicacion"] + list(values_to_save.keys())
                    conn.execute(
                        f"""
                        INSERT INTO comunicaciones ({', '.join(columns)})
                        VALUES ({', '.join(['?'] * len(columns))})
                        """,
                        [communication_id] + list(values_to_save.values()),
                    )
                    audit_change(
                        conn,
                        case_id,
                        "NUEVA COMUNICACIÓN",
                        st.session_state.current_user,
                        new_data=values_to_save,
                        details=description,
                    )

                case_changes = {}
                if values_to_save["es_acuse_recibo"]:
                    case_changes.update({
                        "fecha_acuse": comm_date.isoformat(),
                        "medio_acuse": medium,
                        "resultado_acuse": result_contact,
                        "realizado_acuse_por": performed_by,
                    })
                if values_to_save["es_comunicacion_cierre"]:
                    case_changes.update({
                        "fecha_comunicacion_cierre": comm_date.isoformat(),
                        "medio_comunicacion_cierre": medium,
                        "comunicado_por": performed_by,
                        "resultado_comunicacion_cierre": result_contact,
                    })
                if case_changes:
                    update_case_fields(
                        conn,
                        case_id,
                        case_changes,
                        st.session_state.current_user,
                        "ACTUALIZACIÓN POR COMUNICACIÓN",
                        description,
                    )

            if files:
                case_dir = UPLOAD_DIR / case_id
                case_dir.mkdir(exist_ok=True)
                with db_connection() as conn:
                    for uploaded in files:
                        safe_name = Path(uploaded.name).name
                        target = case_dir / f"{uuid.uuid4().hex[:8]}_{safe_name}"
                        target.write_bytes(uploaded.getbuffer())
                        conn.execute(
                            """
                            INSERT INTO documentos (
                                id_documento, id_caso, id_comunicacion,
                                tipo_documento, nombre_archivo, ruta_archivo,
                                fecha_carga, cargado_por
                            ) VALUES (?, ?, ?, ?, ?, ?, ?, ?)
                            """,
                            (
                                generate_id("DOC"), case_id, communication_id,
                                "Evidencia de comunicación", safe_name,
                                str(target), now_iso(), st.session_state.current_user,
                            ),
                        )
            st.success("Comunicación guardada correctamente.")
            st.rerun()
        except Exception as exc:
            st.error(str(exc))


def render_approval_form(case_id: str):
    with db_connection() as conn:
        case = get_case(conn, case_id)
    if not case:
        return

    tab1, tab2 = st.tabs(["Aprobación interna", "Cierre formal"])
    with tab1:
        with st.form(f"approval_form_{case_id}"):
            recommendation = st.text_area(
                "Recomendación de cierre",
                value=case["recomendacion_cierre"] or "",
                height=130,
            )
            c1, c2 = st.columns(2)
            recommendation_date = c1.date_input(
                "Fecha de recomendación",
                value=date.fromisoformat(case["fecha_recomendacion_cierre"])
                if case["fecha_recomendacion_cierre"] else date.today(),
            )
            recommended_by = c2.text_input(
                "Recomendada por",
                value=case["recomendada_por"] or case["responsable_principal"] or "",
            )
            decision_values = [
                "", "Aprobada", "Devuelta para corrección",
                "Actuación adicional", "Remitida a asesoría jurídica",
            ]
            c1, c2 = st.columns(2)
            decision = c1.selectbox(
                "Decisión del supervisor",
                decision_values,
                index=decision_values.index(case["decision_supervisor"])
                if case["decision_supervisor"] in decision_values else 0,
            )
            decision_date = c2.date_input(
                "Fecha de decisión",
                value=date.fromisoformat(case["fecha_decision_supervisor"])
                if case["fecha_decision_supervisor"] else date.today(),
            )
            supervisor_observation = st.text_area(
                "Observaciones del supervisor",
                value=case["observacion_supervisor"] or "",
            )
            save_approval = st.form_submit_button(
                "Guardar aprobación",
                type="primary",
                use_container_width=True,
            )

        if save_approval:
            try:
                required_text(recommendation, "Recomendación")
                required_text(recommended_by, "Recomendada por")
                new_state = case["estado_actual"]
                if decision == "Aprobada":
                    new_state = "Cierre aprobado"
                elif decision in ("Devuelta para corrección", "Actuación adicional"):
                    new_state = "Devuelta para atención"
                elif decision == "Remitida a asesoría jurídica":
                    new_state = "En atención"

                changes = {
                    "recomendacion_cierre": recommendation,
                    "fecha_recomendacion_cierre": recommendation_date.isoformat(),
                    "recomendada_por": recommended_by,
                    "decision_supervisor": decision,
                    "observacion_supervisor": supervisor_observation,
                    "fecha_decision_supervisor": decision_date.isoformat() if decision else None,
                    "estado_actual": new_state,
                    "situacion_atencion": "En asesoría jurídica"
                    if decision == "Remitida a asesoría jurídica"
                    else case["situacion_atencion"],
                }
                with db_connection() as conn:
                    update_case_fields(
                        conn, case_id, changes,
                        st.session_state.current_user,
                        "RECOMENDACIÓN Y DECISIÓN",
                        supervisor_observation,
                    )
                    register_state_change(
                        conn, case_id, case["estado_actual"], new_state,
                        st.session_state.current_user,
                        supervisor_observation or recommendation,
                    )
                st.success("Aprobación guardada.")
                st.rerun()
            except Exception as exc:
                st.error(str(exc))

    with tab2:
        with st.form(f"closure_form_{case_id}"):
            final_answer = st.text_area(
                "Respuesta final brindada",
                value=case["respuesta_final"] or "",
                height=130,
            )
            c1, c2 = st.columns(2)
            accepted_values = ["", "Sí", "No", "No fue posible contactar", "No aplica"]
            accepted = c1.selectbox(
                "¿Acepta la respuesta?",
                accepted_values,
                index=accepted_values.index(case["acepta_respuesta"])
                if case["acepta_respuesta"] in accepted_values else 0,
            )
            satisfaction_values = [
                "", "Nada satisfecho", "Poco satisfecho",
                "Neutral / Indiferente", "Satisfecho",
                "Muy satisfecho", "No aplica",
            ]
            satisfaction = c2.selectbox(
                "Nivel de satisfacción",
                satisfaction_values,
                index=satisfaction_values.index(case["nivel_satisfaccion"])
                if case["nivel_satisfaccion"] in satisfaction_values else 0,
            )
            dissatisfaction = st.text_area(
                "Comentario cuando la satisfacción sea baja",
                value=case["comentario_insatisfaccion"] or "",
            )
            final_comments = st.text_area(
                "Comentarios finales",
                value=case["comentarios_finales"] or "",
            )
            confidential_close = st.checkbox(
                "La persona solicita que su identidad se mantenga confidencial",
                value=bool(case["confidencial"]),
            )
            c1, c2, c3 = st.columns(3)
            closure_date = c1.date_input(
                "Fecha de cierre",
                value=date.fromisoformat(case["fecha_cierre"])
                if case["fecha_cierre"] else date.today(),
            )
            signed = c2.checkbox(
                "Firma de conformidad",
                value=bool(case["firma_conformidad"]),
            )
            refusal = c2.checkbox(
                "Negativa a firmar",
                value=bool(case["negativa_firma"]),
            )
            witness = st.text_input(
                "Testigo de la comunicación",
                value=case["testigo_cierre"] or "",
            )
            closure_reason = st.text_area(
                "Motivo de cierre",
                value=case["motivo_cierre"] or "",
            )
            files = st.file_uploader(
                "Documento firmado o constancia de cierre",
                accept_multiple_files=True,
                key=f"closure_files_{case_id}",
            )
            save_close = st.form_submit_button(
                "Guardar cierre",
                type="primary",
                use_container_width=True,
            )

        if save_close:
            try:
                required_text(final_answer, "Respuesta final")
                required_text(closure_reason, "Motivo de cierre")
                validate_case_dates(
                    date.fromisoformat(case["fecha_recepcion"]),
                    date.fromisoformat(case["fecha_registro"]),
                    fecha_cierre=closure_date,
                )
                if case["decision_supervisor"] != "Aprobada":
                    raise ValueError("El supervisor debe aprobar el cierre antes de cerrarlo.")
                if case["clasificacion"] == "Queja":
                    if not signed and not refusal:
                        raise ValueError(
                            "Registra firma de conformidad o negativa a firmar."
                        )
                    if refusal:
                        required_text(witness, "Testigo")
                if satisfaction in ("Nada satisfecho", "Poco satisfecho"):
                    required_text(dissatisfaction, "Comentario de insatisfacción")

                changes = {
                    "respuesta_final": final_answer,
                    "acepta_respuesta": accepted,
                    "nivel_satisfaccion": satisfaction,
                    "comentario_insatisfaccion": dissatisfaction,
                    "comentarios_finales": final_comments,
                    "confidencial": int(confidential_close),
                    "fecha_cierre": closure_date.isoformat(),
                    "cerrado_por": st.session_state.current_user,
                    "motivo_cierre": closure_reason,
                    "firma_conformidad": int(signed),
                    "negativa_firma": int(refusal),
                    "testigo_cierre": witness,
                    "estado_actual": "Cerrada",
                }
                with db_connection() as conn:
                    update_case_fields(
                        conn, case_id, changes,
                        st.session_state.current_user,
                        "CIERRE FORMAL", closure_reason,
                    )
                    register_state_change(
                        conn, case_id, case["estado_actual"], "Cerrada",
                        st.session_state.current_user, closure_reason,
                    )
                save_uploaded_files(
                    case_id, files, "Documento de cierre",
                    st.session_state.current_user,
                )
                st.success("Caso cerrado correctamente.")
                st.rerun()
            except Exception as exc:
                st.error(str(exc))


def render_review_form(edit_id: Optional[str] = None):
    existing = None
    if edit_id:
        with db_connection() as conn:
            existing = conn.execute(
                "SELECT * FROM revisiones WHERE id_revision = ?",
                (edit_id,),
            ).fetchone()

    case_id = existing["id_caso"] if existing else select_case_id(
        "Caso asociado", f"review_case_{edit_id or 'new'}"
    )
    if not case_id:
        return
    with db_connection() as conn:
        case = get_case(conn, case_id)
    values = dict(existing) if existing else {}

    with st.form(f"review_form_{edit_id or 'new'}"):
        c1, c2 = st.columns(2)
        type_values = ["Revisión", "Apelación"]
        type_current = values.get("tipo_revision") or "Revisión"
        review_type = c1.selectbox(
            "Tipo",
            type_values,
            index=type_values.index(type_current) if type_current in type_values else 0,
        )
        request_date = c2.date_input(
            "Fecha de solicitud",
            value=date.fromisoformat(values["fecha_solicitud"])
            if values.get("fecha_solicitud") else date.today(),
        )
        requested_by = st.text_input(
            "Solicitada por",
            value=values.get("solicitada_por") or case["nombre_solicitante"] or "Solicitante",
        )
        reason = st.text_area("Motivo de la solicitud", value=values.get("motivo") or "")
        st.caption(
            "La revisión debe ser atendida por el Gerente de la Unidad de Gestión "
            "Ambiental y Social o por una instancia superior independiente designada por la ACP."
        )
        c1, c2 = st.columns(2)
        reviewer = c1.text_input("Persona revisora independiente", value=values.get("revisor") or "")
        status_values = ["Pendiente", "En análisis", "Resuelta"]
        status_current = values.get("estado_revision") or "Pendiente"
        review_status = c2.selectbox(
            "Estado de la revisión",
            status_values,
            index=status_values.index(status_current) if status_current in status_values else 0,
        )
        review_date = st.date_input(
            "Fecha de revisión",
            value=date.fromisoformat(values["fecha_revision"])
            if values.get("fecha_revision") else None,
        )
        decision = st.text_area("Decisión", value=values.get("decision") or "")
        observations = st.text_area("Observaciones", value=values.get("observaciones") or "")
        save = st.form_submit_button(
            "Guardar revisión",
            type="primary",
            use_container_width=True,
        )

    if save:
        try:
            required_text(reason, "Motivo")
            values_to_save = {
                "id_caso": case_id,
                "tipo_revision": review_type,
                "fecha_solicitud": request_date.isoformat(),
                "solicitada_por": requested_by,
                "motivo": reason,
                "revisor": reviewer,
                "fecha_revision": review_date.isoformat() if review_date else None,
                "decision": decision,
                "observaciones": observations,
                "estado_revision": review_status,
            }
            with db_connection() as conn:
                if edit_id:
                    _update(
                        conn, "revisiones", "id_revision", edit_id, values_to_save
                    )
                    audit_change(
                        conn, case_id, "EDICIÓN DE REVISIÓN",
                        st.session_state.current_user,
                        previous_data=dict(existing),
                        new_data=values_to_save,
                        details=reason,
                    )
                else:
                    review_id = generate_id("REV")
                    columns = ["id_revision"] + list(values_to_save.keys())
                    conn.execute(
                        f"""
                        INSERT INTO revisiones ({', '.join(columns)})
                        VALUES ({', '.join(['?'] * len(columns))})
                        """,
                        [review_id] + list(values_to_save.values()),
                    )
                    audit_change(
                        conn, case_id, "REGISTRO DE REVISIÓN",
                        st.session_state.current_user,
                        new_data=values_to_save,
                        details=reason,
                    )

                if review_status != "Resuelta":
                    previous_state = case["estado_actual"]
                    update_case_fields(
                        conn, case_id, {"estado_actual": "En revisión"},
                        st.session_state.current_user,
                        "APERTURA DE REVISIÓN", reason,
                    )
                    register_state_change(
                        conn, case_id, previous_state, "En revisión",
                        st.session_state.current_user, reason,
                    )
            st.success("Revisión guardada correctamente.")
            st.rerun()
        except Exception as exc:
            st.error(str(exc))


def render_traceability():
    case_id = select_case_id("Seleccione el expediente", "trace_case")
    if not case_id:
        return

    with db_connection() as conn:
        case = get_case(conn, case_id)
        states = rows_to_df(conn.execute(
            """
            SELECT fecha_cambio AS fecha, 'Cambio de estado' AS tipo,
                   usuario_cambio AS responsable,
                   COALESCE(estado_anterior, 'Inicio') || ' → ' || estado_nuevo AS evento,
                   motivo AS detalle
            FROM historial_estados WHERE id_caso = ?
            """,
            (case_id,),
        ).fetchall())
        followups = rows_to_df(conn.execute(
            """
            SELECT fecha_registro_sistema AS fecha, 'Seguimiento' AS tipo,
                   responsable_ejecutor AS responsable,
                   tipo_actuacion AS evento,
                   CASE
                       WHEN informo_solicitante = 1 AND COALESCE(contenido_comunicado, '') <> ''
                       THEN descripcion || ' | Información a la persona: ' || contenido_comunicado
                       ELSE descripcion
                   END AS detalle
            FROM seguimientos WHERE id_caso = ?
            """,
            (case_id,),
        ).fetchall())
        documents = rows_to_df(conn.execute(
            """
            SELECT fecha_carga AS fecha, 'Documento' AS tipo,
                   cargado_por AS responsable,
                   tipo_documento AS evento,
                   nombre_archivo AS detalle
            FROM documentos WHERE id_caso = ?
            """,
            (case_id,),
        ).fetchall())
        audits = rows_to_df(conn.execute(
            """
            SELECT fecha_hora AS fecha, 'Auditoría' AS tipo,
                   usuario AS responsable, accion AS evento, detalle
            FROM auditoria WHERE id_caso = ?
            """,
            (case_id,),
        ).fetchall())

    timeline_parts = [
        df for df in [states, followups, documents, audits] if not df.empty
    ]
    timeline = pd.concat(timeline_parts, ignore_index=True) if timeline_parts else pd.DataFrame(
        columns=["fecha", "tipo", "responsable", "evento", "detalle"]
    )
    if not timeline.empty:
        timeline["fecha_orden"] = pd.to_datetime(timeline["fecha"], errors="coerce")
        timeline = timeline.sort_values("fecha_orden", ascending=False).drop(columns=["fecha_orden"])

    c1, c2, c3, c4 = st.columns(4)
    c1.metric("Estado actual", case["estado_actual"])
    c2.metric("Responsable", case["responsable_principal"] or "Sin asignar")
    c3.metric("Seguimientos", len(followups))
    c4.metric("Documentos", len(documents))

    st.markdown("#### Línea de tiempo automática")
    st.caption(
        "Se construye con los cambios de estado, seguimientos, documentos y modificaciones "
        "registradas en el expediente."
    )
    st.dataframe(
        timeline,
        use_container_width=True,
        hide_index=True,
        column_config={"detalle": st.column_config.TextColumn(width="large")},
    )
    render_export_buttons(timeline, [case_id], "traceability", "Trazabilidad")

# =========================================================
# ATENCIÓN Y CIERRE SEGÚN EL INSTRUCTIVO
# =========================================================

RESULTADOS_ATENCION = [
    "En investigación o atención",
    "Atendida sin acción adicional",
    "Atendida con acción adicional",
    "Remitida a otra oficina de la ACP",
    "Remitida a asesoría jurídica",
    "Fuera del objetivo o alcance del mecanismo",
    "Cierre por trámite jurídico mayor de seis meses",
]

ATENDIDA_POR = [
    "ACP - Proyectos Hídricos",
    "Otra oficina de la ACP",
    "Contratista o subcontratista",
]


def render_attention_history():
    followups = followups_history_df()
    communications = communications_history_df()

    if not communications.empty:
        communications = communications.rename(
            columns={
                "id_comunicacion": "id_evento",
                "tipo_comunicacion": "tipo_evento",
                "fecha": "fecha_evento",
                "realizada_por": "responsable_evento",
                "resultado_contacto": "resultado_evento",
            }
        )
        communications["origen_evento"] = "Información al solicitante"
    if not followups.empty:
        followups = followups.rename(
            columns={
                "id_seguimiento": "id_evento",
                "tipo_actuacion": "tipo_evento",
                "fecha_actuacion": "fecha_evento",
                "responsable_ejecutor": "responsable_evento",
                "resultado": "resultado_evento",
            }
        )
        followups["origen_evento"] = "Seguimiento"

    common = [
        "id_evento", "id_caso", "codigo_caso", "clasificacion",
        "origen_evento", "fecha_evento", "tipo_evento",
        "descripcion", "responsable_evento", "resultado_evento",
    ]
    parts = []
    if not followups.empty:
        parts.append(followups[[c for c in common if c in followups.columns]])
    if not communications.empty:
        parts.append(communications[[c for c in common if c in communications.columns]])

    df = pd.concat(parts, ignore_index=True) if parts else pd.DataFrame(columns=common)
    df = filter_dataframe_ui(df, "attention_history")
    if not df.empty:
        df = df.sort_values("fecha_evento", ascending=False)

    selected_records = dataframe_selection(
        df,
        "attention_history",
        "id_evento",
        [
            "codigo_caso", "clasificacion", "origen_evento",
            "fecha_evento", "tipo_evento", "descripcion",
            "responsable_evento", "resultado_evento",
        ],
    )
    selected_case_ids = []
    if selected_records and not df.empty:
        selected_case_ids = (
            df[df["id_evento"].astype(str).isin(selected_records)]["id_caso"]
            .dropna().astype(str).unique().tolist()
        )
    render_export_buttons(
        df,
        selected_case_ids,
        "attention_history",
        "Atencion_y_seguimiento",
    )


def render_attention_form():
    case_id = select_case_id("Seleccione el caso", "attention_case")
    if not case_id:
        return
    with db_connection() as conn:
        case = get_case(conn, case_id)

    st.caption(
        "Registre en una sola actuación la gestión realizada y, cuando corresponda, "
        "la información proporcionada a la persona."
    )

    with st.form(f"attention_form_{case_id}"):
        c1, c2 = st.columns(2)
        action_date = c1.date_input("Fecha de la actuación", value=date.today())
        action_time = time_select_12h(
            c2, "Hora de la actuación", None, key=f"attention_time_{case_id}"
        )
        c1, c2 = st.columns(2)
        action_type = c1.selectbox("Tipo de actuación", FOLLOWUP_TYPES)
        executor = c2.text_input(
            "Responsable de atender",
            value=case["responsable_principal"] or "",
        )
        description = st.text_area(
            "Descripción del seguimiento del caso",
            height=160,
        )
        result = st.text_area("Resultado obtenido", height=130)

        c1, c2 = st.columns(2)
        next_action = c1.text_area("Próxima acción", height=110)
        commitment_date = c2.date_input("Fecha compromiso", value=None)

        st.markdown("#### Información proporcionada a la persona")
        informed = st.checkbox(
            "Se informó a la persona sobre la recepción, avance o resultado"
        )
        c1, c2, c3 = st.columns(3)
        information_type_options = _options_with_current(COMMUNICATION_TYPES, "", False)
        communication_type = c1.selectbox(
            "Tipo de información",
            information_type_options,
        )
        communication_medium = c2.selectbox(
            "Medio utilizado",
            MEDIOS_RECEPCION,
        )
        communication_result = c3.selectbox(
            "Resultado del contacto",
            [
                "Recibido",
                "No localizado",
                "Sin respuesta",
                "Datos de contacto no disponibles",
                "Rechazó la comunicación",
                "Otro",
            ],
        )
        communication_text = st.text_area(
            "Contenido comunicado",
            height=130,
        )

        files = st.file_uploader(
            "Informes, evidencias o adjuntos",
            accept_multiple_files=True,
            key=f"attention_files_{case_id}",
        )
        save = st.form_submit_button(
            "Guardar seguimiento",
            type="primary",
            use_container_width=True,
        )

    if save:
        try:
            required_text(executor, "Responsable de atender")
            required_text(description, "Descripción del seguimiento")
            if informed:
                required_text(communication_text, "Contenido comunicado")
            if action_date < date.fromisoformat(case["fecha_registro"]):
                raise ValueError("La actuación no puede ser anterior al registro del caso.")
            if commitment_date and commitment_date < action_date:
                raise ValueError("La fecha compromiso no puede ser anterior a la actuación.")

            followup_id = generate_id("SEG")
            values_to_save = {
                "id_seguimiento": followup_id,
                "id_caso": case_id,
                "fecha_actuacion": action_date.isoformat(),
                "hora_actuacion": action_time,
                "tipo_actuacion": action_type,
                "descripcion": normalize_text(description),
                "resultado": normalize_text(result),
                "responsable_ejecutor": normalize_text(executor),
                "usuario_registro": st.session_state.current_user,
                "estado_anterior": case["estado_actual"],
                "estado_posterior": "En atención",
                "proxima_accion": normalize_text(next_action),
                "fecha_compromiso": commitment_date.isoformat() if commitment_date else None,
                "estado_actividad": "Completada",
                "visible_solicitante": int(informed),
                "informo_solicitante": int(informed),
                "tipo_informacion_solicitante": communication_type if informed else "",
                "medio_informacion_solicitante": communication_medium if informed else "",
                "resultado_contacto_solicitante": communication_result if informed else "",
                "contenido_comunicado": normalize_text(communication_text) if informed else "",
                "fecha_registro_sistema": now_iso(),
            }

            with db_connection() as conn:
                columns = list(values_to_save.keys())
                conn.execute(
                    f"INSERT INTO seguimientos ({', '.join(columns)}) "
                    f"VALUES ({', '.join(['?'] * len(columns))})",
                    list(values_to_save.values()),
                )
                _assign_control_code(
                    conn, "seguimientos", "id_seguimiento", followup_id,
                    "id_control_seguimiento", "M8-S-"
                )

                if case["estado_actual"] != "En atención":
                    update_case_fields(
                        conn,
                        case_id,
                        {"estado_actual": "En atención"},
                        st.session_state.current_user,
                        "ACTUALIZACIÓN POR SEGUIMIENTO",
                        description,
                    )
                    register_state_change(
                        conn,
                        case_id,
                        case["estado_actual"],
                        "En atención",
                        st.session_state.current_user,
                        description,
                        followup_id,
                    )

                audit_change(
                    conn,
                    case_id,
                    "SEGUIMIENTO DEL CASO",
                    st.session_state.current_user,
                    new_data=values_to_save,
                    details=description,
                )

            if files:
                case_dir = UPLOAD_DIR / case_id
                case_dir.mkdir(exist_ok=True)
                with db_connection() as conn:
                    for uploaded in files:
                        safe_name = Path(uploaded.name).name
                        target = case_dir / f"{uuid.uuid4().hex[:8]}_{safe_name}"
                        target.write_bytes(uploaded.getbuffer())
                        conn.execute(
                            """
                            INSERT INTO documentos (
                                id_documento, id_caso, id_seguimiento,
                                tipo_documento, nombre_archivo,
                                ruta_archivo, fecha_carga, cargado_por
                            ) VALUES (?, ?, ?, ?, ?, ?, ?, ?)
                            """,
                            (
                                generate_id("DOC"), case_id, followup_id,
                                "Evidencia de seguimiento", safe_name,
                                str(target), now_iso(), st.session_state.current_user,
                            ),
                        )
            st.success("Seguimiento guardado correctamente.")
            st.rerun()
        except Exception as exc:
            st.error(str(exc))

def render_closure_flow():
    case_id = select_case_id("Seleccione el caso", "closure_case_v11")
    if not case_id:
        return

    with db_connection() as conn:
        case_row = get_case(conn, case_id)
    if not case_row:
        return
    case = dict(case_row)

    st.caption(
        "El cierre se gestiona en dos etapas: preparación y visto bueno; después, "
        "comunicación y cierre final. Las investigaciones, remisiones y demás actuaciones "
        "se documentan en Seguimiento."
    )

    st.markdown("#### Resumen del expediente")
    c1, c2, c3, c4 = st.columns(4)
    c1.metric("Código", case.get("codigo_caso") or "")
    c2.metric("Clasificación", case.get("clasificacion") or "")
    c3.metric("Estado", case.get("estado_actual") or "")
    c4.metric("Responsable", case.get("responsable_principal") or "Sin asignar")

    if case.get("clasificacion") == "Queja":
        try:
            validate_required_case_location(case)
            st.success("La ubicación obligatoria del caso está completa.")
        except ValueError as exc:
            st.warning(str(exc))

    st.markdown("### 1. Respuesta, propuesta y visto bueno")
    with st.form(f"closure_preparation_{case_id}"):
        final_answer = st.text_area(
            "Respuesta brindada al caso",
            value=case.get("respuesta_final") or "",
            height=180,
            help="La respuesta debe resumir claramente el resultado que será comunicado a la persona.",
        )
        recommendation = st.text_area(
            "Resumen y fundamento de la propuesta de cierre",
            value=case.get("recomendacion_cierre") or "",
            height=150,
            help="Campo interno requerido por el instructivo para sustentar el cierre.",
        )
        c1, c2 = st.columns(2)
        recommendation_date = c1.date_input(
            "Fecha de la propuesta de cierre",
            value=date.fromisoformat(case["fecha_recomendacion_cierre"])
            if case.get("fecha_recomendacion_cierre") else date.today(),
        )
        recommended_by = c2.text_input(
            "Responsable que propone el cierre",
            value=case.get("recomendada_por") or case.get("responsable_principal") or "",
        )

        current_decision = case.get("decision_supervisor") or "Pendiente"
        decision_aliases = {
            "Aprobada": "Aprobado",
            "Sí": "Aprobado",
            "No aprobada": "Devuelto para seguimiento",
            "No": "Devuelto para seguimiento",
        }
        current_decision = decision_aliases.get(current_decision, current_decision)
        decision_options = ["Pendiente", "Aprobado", "Devuelto para seguimiento"]
        if current_decision not in decision_options:
            decision_options.append(current_decision)
        c1, c2 = st.columns(2)
        supervisor_decision = c1.selectbox(
            "Visto bueno del supervisor",
            decision_options,
            index=decision_options.index(current_decision),
        )
        supervisor_date = c2.date_input(
            "Fecha de la decisión del supervisor",
            value=date.fromisoformat(case["fecha_decision_supervisor"])
            if case.get("fecha_decision_supervisor") else None,
        )
        supervisor_observation = st.text_area(
            "Observaciones del supervisor",
            value=case.get("observacion_supervisor")
            or case.get("observacion_visto_bueno")
            or "",
            height=130,
        )
        save_preparation = st.form_submit_button(
            "Guardar respuesta y visto bueno",
            type="primary",
            use_container_width=True,
        )

    if save_preparation:
        try:
            required_text(final_answer, "Respuesta brindada al caso")
            required_text(recommendation, "Resumen y fundamento de la propuesta de cierre")
            required_text(recommended_by, "Responsable que propone el cierre")
            if supervisor_decision != "Pendiente" and not supervisor_date:
                raise ValueError("Registre la fecha de la decisión del supervisor.")
            if supervisor_decision == "Devuelto para seguimiento":
                required_text(supervisor_observation, "Observaciones del supervisor")

            new_state = {
                "Pendiente": "Pendiente de aprobación",
                "Aprobado": "Cierre aprobado",
                "Devuelto para seguimiento": "En atención",
            }[supervisor_decision]
            changes = {
                "respuesta_final": normalize_text(final_answer),
                "se_brindo_respuesta": "Sí",
                "recomendacion_cierre": normalize_text(recommendation),
                "fecha_recomendacion_cierre": recommendation_date.isoformat(),
                "recomendada_por": normalize_text(recommended_by),
                "decision_supervisor": supervisor_decision,
                "visto_bueno_supervisor": (
                    "Sí" if supervisor_decision == "Aprobado"
                    else "No" if supervisor_decision == "Devuelto para seguimiento"
                    else "Pendiente"
                ),
                "fecha_visto_bueno": supervisor_date.isoformat() if supervisor_date else None,
                "fecha_decision_supervisor": supervisor_date.isoformat() if supervisor_date else None,
                "observacion_visto_bueno": normalize_text(supervisor_observation),
                "observacion_supervisor": normalize_text(supervisor_observation),
                "estado_actual": new_state,
            }
            with db_connection() as conn:
                update_case_fields(
                    conn,
                    case_id,
                    changes,
                    st.session_state.current_user,
                    "RESPUESTA Y VISTO BUENO DE CIERRE",
                    supervisor_observation or recommendation,
                )
                register_state_change(
                    conn,
                    case_id,
                    case.get("estado_actual"),
                    new_state,
                    st.session_state.current_user,
                    supervisor_observation or recommendation,
                )
            if supervisor_decision == "Devuelto para seguimiento":
                st.warning("El caso fue devuelto a Seguimiento con las observaciones registradas.")
            else:
                st.success("Respuesta y visto bueno guardados correctamente.")
            st.rerun()
        except Exception as exc:
            st.error(str(exc))

    approved = (
        case.get("decision_supervisor") in {"Aprobado", "Aprobada"}
        or case.get("visto_bueno_supervisor") == "Sí"
    )
    if case.get("estado_actual") == "Cerrada":
        st.success("Este caso ya se encuentra cerrado. Puede consultar su expediente en Histórico.")
        return
    if not approved:
        st.info(
            "La comunicación y el cierre final se habilitan cuando el supervisor aprueba la propuesta."
        )
        return

    st.markdown("### 2. Comunicación y cierre final")
    st.info("Respuesta aprobada: " + (case.get("respuesta_final") or "Sin texto registrado."))

    with st.form(f"closure_final_{case_id}"):
        c1, c2 = st.columns(2)
        communication_date = c1.date_input(
            "Fecha de comunicación del resultado",
            value=date.fromisoformat(case["fecha_comunicacion_cierre"])
            if case.get("fecha_comunicacion_cierre") else date.today(),
        )
        communication_time = time_select_12h(
            c2,
            "Hora de comunicación del resultado",
            case.get("hora_comunicacion_cierre"),
            key=f"communication_close_time_{case_id}",
        )
        c1, c2 = st.columns(2)
        medium_options = _options_with_current(
            RECEPTION_CHANNELS,
            case.get("medio_comunicacion_cierre"),
            False,
        )
        communication_medium = c1.selectbox(
            "Medio de comunicación",
            medium_options,
            index=_option_index(
                medium_options,
                case.get("medio_comunicacion_cierre") or RECEPTION_CHANNELS[0],
            ),
        )
        communicated_by = c2.text_input(
            "Resultado comunicado por",
            value=case.get("comunicado_por") or st.session_state.current_user,
        )

        c1, c2 = st.columns(2)
        accepted_current = case.get("acepta_respuesta") or "Pendiente"
        accepted_options = _options_with_current(
            ACEPTACION_RESPUESTA_OPTIONS, accepted_current, False
        )
        accepted = c1.selectbox(
            "¿La persona acepta la respuesta?",
            accepted_options,
            index=_option_index(accepted_options, accepted_current),
        )
        satisfaction_current = case.get("nivel_satisfaccion") or ""
        satisfaction_options = _options_with_current(
            SATISFACCION_OPTIONS, satisfaction_current
        )
        satisfaction = c2.selectbox(
            "Nivel de satisfacción",
            satisfaction_options,
            index=_option_index(satisfaction_options, satisfaction_current),
        )
        comments_acceptance = st.text_area(
            "Comentarios sobre la aceptación de la respuesta",
            value=case.get("comentarios_aceptacion_contacto") or "",
            height=120,
        )
        comments_rejection = st.text_area(
            "Comentarios sobre el rechazo o la insatisfacción",
            value=case.get("comentarios_rechazo_contacto")
            or case.get("comentario_insatisfaccion")
            or "",
            height=120,
        )
        disclosure_current = case.get("autoriza_divulgacion_datos") or (
            "No" if bool(case.get("confidencial")) else "Sí"
        )
        authorizes_disclosure = st.selectbox(
            "¿Autoriza compartir sus datos personales en reportes o divulgaciones?",
            _options_with_current(SI_NO_OPTIONS, disclosure_current, False),
            index=_option_index(
                _options_with_current(SI_NO_OPTIONS, disclosure_current, False),
                disclosure_current,
            ),
        )
        final_comments = st.text_area(
            "Comentarios finales",
            value=case.get("comentarios_finales") or "",
            height=130,
        )

        c1, c2 = st.columns(2)
        closure_date = c1.date_input(
            "Fecha de cierre",
            value=date.fromisoformat(case["fecha_cierre"])
            if case.get("fecha_cierre") else date.today(),
        )
        closure_time = time_select_12h(
            c2,
            "Hora de cierre",
            case.get("hora_cierre"),
            key=f"closure_time_v11_{case_id}",
        )
        c1, c2 = st.columns(2)
        signed = c1.checkbox(
            "Firma de conformidad",
            value=bool(case.get("firma_conformidad")),
        )
        refusal = c2.checkbox(
            "La persona se negó a firmar",
            value=bool(case.get("negativa_firma")),
        )
        witness = st.text_input(
            "Testigo de la comunicación, cuando exista negativa a firmar",
            value=case.get("testigo_cierre") or "",
        )
        closure_reason = st.text_area(
            "Observación final del cierre",
            value=case.get("motivo_cierre") or "",
            height=120,
        )
        files = st.file_uploader(
            "Formulario firmado o constancia de cierre",
            accept_multiple_files=True,
            key=f"closure_final_files_{case_id}",
        )
        save_close = st.form_submit_button(
            "Cerrar caso",
            type="primary",
            use_container_width=True,
        )

    if save_close:
        try:
            validate_required_case_location(case)
            required_text(case.get("respuesta_final") or "", "Respuesta brindada al caso")
            required_text(communicated_by, "Resultado comunicado por")
            if accepted not in ("Sí", "No"):
                raise ValueError("Registre si la persona acepta o no la respuesta.")
            required_text(satisfaction, "Nivel de satisfacción")
            if accepted == "No":
                required_text(comments_rejection, "Comentario sobre el rechazo")
            if satisfaction in ("1 - Nada satisfecho", "2 - Poco satisfecho"):
                required_text(comments_rejection, "Comentario de satisfacción baja")
            if case.get("clasificacion") == "Queja":
                if not signed and not refusal:
                    raise ValueError(
                        "Por estar clasificado como Queja, el caso requiere firma de conformidad "
                        "o constancia de negativa a firmar."
                    )
                if refusal:
                    required_text(witness, "Testigo de la comunicación")
            validate_case_dates(
                date.fromisoformat(case["fecha_recepcion"]),
                date.fromisoformat(case["fecha_registro"]),
                fecha_cierre=closure_date,
            )

            changes = {
                "fecha_comunicacion_cierre": communication_date.isoformat(),
                "hora_comunicacion_cierre": communication_time,
                "medio_comunicacion_cierre": communication_medium,
                "comunicado_por": normalize_text(communicated_by),
                "resultado_comunicacion_cierre": accepted,
                "acepta_respuesta": accepted,
                "nivel_satisfaccion": satisfaction,
                "comentarios_aceptacion_contacto": normalize_text(comments_acceptance),
                "comentarios_rechazo_contacto": normalize_text(comments_rejection),
                "comentario_insatisfaccion": normalize_text(comments_rejection),
                "autoriza_divulgacion_datos": authorizes_disclosure,
                "confidencial": int(authorizes_disclosure == "No"),
                "comentarios_finales": normalize_text(final_comments),
                "fecha_cierre": closure_date.isoformat(),
                "hora_cierre": closure_time,
                "cerrado_por": st.session_state.current_user,
                "motivo_cierre": normalize_text(closure_reason),
                "firma_conformidad": int(signed),
                "negativa_firma": int(refusal),
                "testigo_cierre": normalize_text(witness),
                "estado_actual": "Cerrada",
            }

            with db_connection() as conn:
                # Toda comunicación se incorpora a Seguimiento, no a una pantalla separada.
                followup_id = generate_id("SEG")
                communication_description = (
                    f"Se comunicó la respuesta final del caso por {communication_medium}. "
                    f"Aceptación registrada: {accepted}. Nivel de satisfacción: {satisfaction}."
                )
                conn.execute(
                    """
                    INSERT INTO seguimientos (
                        id_seguimiento, id_caso, fecha_actuacion, hora_actuacion,
                        tipo_actuacion, descripcion, resultado,
                        responsable_ejecutor, usuario_registro,
                        estado_anterior, estado_posterior,
                        proxima_accion, fecha_compromiso, estado_actividad,
                        visible_solicitante, fecha_registro_sistema,
                        informo_solicitante, tipo_informacion_solicitante,
                        medio_informacion_solicitante, resultado_contacto_solicitante,
                        contenido_comunicado
                    ) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
                    """,
                    (
                        followup_id, case_id, communication_date.isoformat(), communication_time,
                        "Comunicación de resultado y cierre", communication_description,
                        case.get("respuesta_final") or "", normalize_text(communicated_by),
                        st.session_state.current_user, case.get("estado_actual"), "Cerrada",
                        "", None, "Completada", 1, now_iso(), 1,
                        "Comunicación de resultado", communication_medium, accepted,
                        case.get("respuesta_final") or "",
                    ),
                )
                _assign_control_code(
                    conn,
                    "seguimientos",
                    "id_seguimiento",
                    followup_id,
                    "id_control_seguimiento",
                    "M8-S-",
                )
                update_case_fields(
                    conn,
                    case_id,
                    changes,
                    st.session_state.current_user,
                    "CIERRE FINAL DEL CASO",
                    closure_reason or "Comunicación y cierre final",
                )
                register_state_change(
                    conn,
                    case_id,
                    case.get("estado_actual"),
                    "Cerrada",
                    st.session_state.current_user,
                    closure_reason or "Comunicación y cierre final",
                    followup_id,
                )
                audit_change(
                    conn,
                    case_id,
                    "COMUNICACIÓN INCORPORADA AL SEGUIMIENTO",
                    st.session_state.current_user,
                    new_data={"id_seguimiento": followup_id},
                    details=communication_description,
                )

            save_uploaded_files(
                case_id,
                files,
                "Documento de cierre",
                st.session_state.current_user,
            )
            st.success("Caso cerrado correctamente.")
            st.rerun()
        except Exception as exc:
            st.error(str(exc))

# =========================================================
# PANTALLAS
# =========================================================

if page == "Panel general":
    mostrar_ayuda_pantalla("Panel general")
    with db_connection() as conn:
        total = conn.execute("SELECT COUNT(*) total FROM casos").fetchone()["total"]
        active = conn.execute(
            "SELECT COUNT(*) total FROM casos WHERE estado_actual NOT IN ('Cerrada', 'Fuera de alcance')"
        ).fetchone()["total"]
        closed = conn.execute(
            "SELECT COUNT(*) total FROM casos WHERE estado_actual = 'Cerrada'"
        ).fetchone()["total"]
        linked = conn.execute(
            "SELECT COUNT(*) total FROM casos WHERE pertenece_proyecto = 1"
        ).fetchone()["total"]
        external = conn.execute(
            "SELECT COUNT(*) total FROM casos WHERE pertenece_proyecto = 0"
        ).fetchone()["total"]
        overdue_ack = conn.execute(
            """
            SELECT COUNT(*) total FROM casos
            WHERE fecha_acuse IS NULL
              AND estado_actual NOT IN ('Cerrada', 'Fuera de alcance')
              AND julianday('now') - julianday(fecha_registro) > 7
            """
        ).fetchone()["total"]

    c1, c2, c3, c4, c5, c6 = st.columns(6)
    c1.metric("Casos registrados", total)
    c2.metric("Casos activos", active)
    c3.metric("Casos cerrados", closed)
    c4.metric("Vinculados con M01", linked)
    c5.metric("Personas externas a M01", external)
    c6.metric("Sin acuse oportuno", overdue_ack)

    st.markdown("### Casos recientes")
    recent = cases_history_df().head(30)
    st.dataframe(
        recent.drop(columns=["id_caso"], errors="ignore"),
        use_container_width=True,
        hide_index=True,
    )
    st.download_button(
        "Descargar todas las tablas en Excel",
        data=export_all_tables_excel(),
        file_name="modulo_8_casos_todas_las_tablas.xlsx",
        mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        use_container_width=True,
    )

elif page == "Carga de información":
    st.markdown("### Carga de información desde Survey123")
    mostrar_ayuda_pantalla("Carga de información")
    uploaded_file = st.file_uploader(
        "Cargar archivo Excel exportado desde Survey123",
        type=["xlsx"],
        key="survey_import_v11",
    )
    if uploaded_file:
        try:
            uploaded_file.seek(0)
            preview = pd.ExcelFile(uploaded_file)
            st.write("Hojas detectadas:", ", ".join(preview.sheet_names))
            structure = survey_structure_report(uploaded_file)
            st.markdown("#### Validación de estructura")
            st.dataframe(structure, use_container_width=True, hide_index=True)
            st.caption(
                "Los IDs internos M8 se generan automáticamente. GlobalID y ParentGlobalID "
                "se conservan como referencias para el intercambio futuro con el cliente."
            )
            if st.button(
                "Importar o actualizar Survey123",
                type="primary",
                use_container_width=True,
            ):
                uploaded_file.seek(0)
                result = importar_survey123(
                    uploaded_file, st.session_state.current_user
                )
                st.success("Importación Survey123 finalizada.")
                st.json(result)
                st.rerun()
        except Exception as exc:
            st.error(f"No fue posible procesar el archivo: {exc}")

    pending = read_table_df("registros_survey_pendientes")
    if not pending.empty:
        st.markdown("#### Registros relacionados pendientes de vinculación")
        st.caption(
            "Se conservaron completos porque su caso padre no estaba incluido en el archivo."
        )
        st.dataframe(pending, use_container_width=True, hide_index=True)

    history = read_table_df("importaciones_survey123")
    if not history.empty:
        st.markdown("#### Historial de importaciones Survey123")
        st.dataframe(history, use_container_width=True, hide_index=True)

elif page == "Registro de casos":
    st.markdown("### Registro de casos")
    mostrar_ayuda_pantalla("Registro de casos")
    st.caption(
        "La cédula permite vincular el caso con una persona y un hogar de M01 cuando existe "
        "coincidencia. Si no existe, el caso se registra normalmente como persona externa."
    )
    tab_registered, tab_new, tab_edit = st.tabs(
        ["Casos registrados", "Agregar nuevo caso", "Editar caso"]
    )
    with tab_registered:
        render_cases_history()
    with tab_new:
        submitted, payload, files = case_form(form_key="new_case_v11")
        if submitted:
            try:
                code_created = save_new_case(payload, files)
                st.success(f"Caso registrado correctamente: {code_created}")
                st.rerun()
            except Exception as exc:
                st.error(str(exc))
    with tab_edit:
        case_id = select_case_id("Seleccione el caso a editar", "edit_case_v11")
        if case_id:
            with db_connection() as conn:
                selected_case = get_case(conn, case_id)
            submitted, payload, files = case_form(
                defaults=selected_case,
                form_key=f"edit_case_v11_{case_id}",
            )
            if submitted:
                try:
                    update_existing_case(case_id, payload, files)
                    st.success("Caso actualizado correctamente.")
                    st.rerun()
                except Exception as exc:
                    st.error(str(exc))

elif page == "Seguimiento":
    st.markdown("### Seguimiento de casos")
    mostrar_ayuda_pantalla("Seguimiento")
    mode = render_work_mode(
        ["Agregar seguimiento", "Editar seguimiento"],
        "followup_mode_v11",
    )
    if mode == "Agregar seguimiento":
        render_attention_form()
    else:
        df = followups_history_df()
        if df.empty:
            st.info("No hay seguimientos para editar.")
        else:
            labels = {
                f"{row['codigo_caso']} · {row['fecha_actuacion']} · "
                f"{row['tipo_actuacion']} · {str(row['descripcion'])[:70]}": row["id_seguimiento"]
                for _, row in df.iterrows()
            }
            selected = st.selectbox("Seleccione el seguimiento", list(labels.keys()))
            render_followup_form(labels[selected])

elif page == "Cierre":
    st.markdown("### Cierre de casos")
    mostrar_ayuda_pantalla("Cierre")
    render_closure_flow()

elif page == "Histórico":
    st.markdown("### Histórico final del módulo")
    mostrar_ayuda_pantalla("Histórico")
    tab_cases, tab_followups, tab_closures, tab_trace = st.tabs(
        ["Casos", "Seguimientos", "Cierres", "Trazabilidad"]
    )
    with tab_cases:
        render_cases_history()
    with tab_followups:
        render_followups_history()
    with tab_closures:
        render_approvals_history()
    with tab_trace:
        render_traceability()

